"""
MÓDULO DE COMPILACIÓN DE INFORMES DOCX Y EXPORTACIÓN A PDF
===========================================================
Autor: Federico Agustín Chillón
Genera los documentos institucionales con diseño y paginación equilibrada:
1. Paper Semanal de Investigación Macroeconómica (APA 7)
2. Informe Mensual de Coyuntura Macroeconómica y Regional (OERU - UNCUYO)
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import win32com.client
import pandas as pd

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def formatear_tabla_apa7(tabla, col_widths, headers, data_rows):
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Encabezado
    hdr_cells = tabla.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Filas de datos
    for r_idx, row_data in enumerate(data_rows):
        row_cells = tabla.add_row().cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if isinstance(val, (int, float)) or ('%' in str(val) or '$' in str(val)) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.0)
                run.font.color.rgb = RGBColor(30, 41, 59)
                
    for row in tabla.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)


def compilar_paper_semanal(ruta_excel: str, dir_figuras: str, ruta_salida_docx: str):
    doc = docx.Document()
    
    # Márgenes de página optimizados
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
    # Título Principal
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Paper de Investigación Macroeconómica y Estrategia Financiera")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(15)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(2)
    
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("Período: Semana del 17 al 21 de Agosto de 2026 | Autor: Federico Agustín Chillón\nMarco Institucional: Análisis de Renta Fija, Microestructura Cambiaria y Régimen Monetario")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(9.0)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(71, 85, 105)
    p_meta.paragraph_format.space_after = Pt(10)
    
    # 1. Resumen Ejecutivo
    h1 = doc.add_heading("1. Resumen Ejecutivo y Diagnóstico del Régimen Macroeconómico", level=2)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(4)
    p1 = doc.add_paragraph(
        "Durante la semana bajo análisis, el mercado financiero argentino transitó una fase de descompresión de tasas en los instrumentos de deuda soberana y estabilidad cambiaria relativa. La política macroeconómica continúa anclada en tres pilares: (i) ancla fiscal basada en el superávit financiero primario del Sector Público Nacional, (ii) saneamiento del balance del Banco Central mediante la transferencia de pasivos remunerados a Letras Fiscales de Liquidez (Lefi), y (iii) deslizamiento administrado (crawling peg) de la cotización oficial minorista."
    )
    p1.paragraph_format.space_after = Pt(8)
    
    # 2. Microestructura Cambiaria
    h2 = doc.add_heading("2. Microestructura Cambiaria, Brechas y Futuros ROFEX", level=2)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(4)
    p2 = doc.add_paragraph(
        "Las cotizaciones cambiarias spot y financieras mostraron un estrechamiento de brechas. El Dólar Oficial Minorista (BNA) finalizó en $1.515,00, mientras que el Contado con Liquidación (CCL) cerró en $1.596,59, situando la brecha cambiaria implícita en 5,39%. En el mercado de derivados Matba-Rofex, las posiciones a 1 y 3 meses operaron con tasas nominales anuales implícitas alineadas a la tasa de política monetaria."
    )
    p2.paragraph_format.space_after = Pt(6)
    
    # Figura Microestructura (Página 1)
    fig3_path = os.path.join(dir_figuras, "Microestructura_Cambiaria_v2.png")
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6.3))
        p_cap = doc.add_paragraph()
        run_cap = p_cap.add_run("Figura 1. Microestructura de tipos de cambio spot, financieros y brecha implícita CCL/Oficial.")
        run_cap.font.size = Pt(8.0)
        run_cap.font.bold = True
        p_cap.paragraph_format.space_after = Pt(6)
        
    # Salto de página explícito para página 2 perfecta
    doc.add_page_break()
    
    # 3. Curvas Soberanas (Página 2)
    h3 = doc.add_heading("3. Estructura Temporal de Tasas de Interés y Curvas Soberanas", level=2)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(4)
    p3 = doc.add_paragraph(
        "La curva soberana en moneda dura (hard dollar) operó con un desplazamiento descendente en las TIRs, ubicándose en el rango 9,70% - 12,60%. El spread por legislación entre bonos Globales (Ley NY) y Bonares (Ley Local) se consolidó entre 40 y 50 puntos básicos. En el tramo en pesos a tasa fija, las LECAPS mantuvieron tasas efectivas mensuales en el rango 2,95% - 3,15%, reflejando una tasa real positiva ex-ante frente a la inflación esperada del REM."
    )
    p3.paragraph_format.space_after = Pt(6)
    
    # Tabla de Bonos Soberanos (Página 2)
    df_usd = pd.read_excel(ruta_excel, sheet_name="Curva_Soberana_USD")
    t_bonos = doc.add_table(rows=1, cols=6)
    formatear_tabla_apa7(
        t_bonos,
        col_widths=[1.0, 1.3, 1.0, 1.0, 1.0, 1.0],
        headers=["Ticker", "Legislación", "Vencimiento", "Precio (USD)", "TIR (%)", "Duration Mod."],
        data_rows=df_usd[['Ticker', 'Legislacion', 'Maturity', 'Precio_USD', 'TIR_%', 'Modified_Duration']].values.tolist()
    )
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)
    
    # Figura Curvas (Página 2)
    fig1_path = os.path.join(dir_figuras, "Curva_Rendimientos_Soberanos_v2.png")
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.3))
        p_cap1 = doc.add_paragraph()
        run_cap1 = p_cap1.add_run("Figura 2. Estructura temporal de rendimientos soberanos en USD y contraste LECAP vs. REM.")
        run_cap1.font.size = Pt(8.0)
        run_cap1.font.bold = True
        p_cap1.paragraph_format.space_after = Pt(4)

    os.makedirs(os.path.dirname(os.path.abspath(ruta_salida_docx)), exist_ok=True)
    doc.save(ruta_salida_docx)
    return ruta_salida_docx


def compilar_informe_oeru(ruta_excel: str, dir_figuras: str, ruta_salida_docx: str):
    doc = docx.Document()
    
    # Márgenes de página
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
    # Título OERU
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Informe Mensual de Coyuntura Macroeconómica y Regional")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(15)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(2)
    
    p_inst = doc.add_paragraph()
    run_inst = p_inst.add_run("Observatorio Económico Regional Urbano (OERU) | Facultad de Ciencias Económicas — UNCUYO\nAutor: Federico Agustín Chillón | Edición: Agosto 2026")
    run_inst.font.name = "Arial"
    run_inst.font.size = Pt(9.0)
    run_inst.font.italic = True
    run_inst.font.color.rgb = RGBColor(71, 85, 105)
    p_inst.paragraph_format.space_after = Pt(10)
    
    # 1. Inflación y Salarios (Página 1)
    h1 = doc.add_heading("1. Dinámica de Precios, Salarios Reales y Canastas Básicas", level=2)
    h1.paragraph_format.space_before = Pt(4)
    h1.paragraph_format.space_after = Pt(4)
    p1 = doc.add_paragraph(
        "El relevamiento de precios al consumidor confirmó la convergencia inflacionaria hacia el nivel de 2,2% mensual a nivel nacional (INDEC) y 2,2% en la provincia de Mendoza (DEIE). La Canasta Básica Total (CBT) en Mendoza alcanzó los $960.000 para un hogar tipo 2. El salario registrado medido por el índice RIPTE exhibió una variación mensual de 2,4%, consolidando una leve recuperación del poder adquisitivo en términos reales."
    )
    p1.paragraph_format.space_after = Pt(6)
    
    # Tabla Inflación y Actividad (Página 1)
    df_inf = pd.read_excel(ruta_excel, sheet_name="Inflacion_Salarios_Actividad")
    t_inf = doc.add_table(rows=1, cols=6)
    formatear_tabla_apa7(
        t_inf,
        col_widths=[1.0, 1.1, 1.1, 1.1, 1.1, 1.1],
        headers=["Mes", "IPC INDEC", "IPC Núcleo", "IPC Mendoza", "EMAE YoY", "Despachos INV"],
        data_rows=df_inf[['Mes', 'IPC_General_Nacional_INDEC_%', 'IPC_Nucleo_INDEC_%', 'IPC_Mendoza_DEIE_%', 'EMAE_Var_Interanual_%', 'Despachos_Vino_INV_Miles_HL']].values.tolist()
    )
    p_space1 = doc.add_paragraph()
    p_space1.paragraph_format.space_after = Pt(6)
    
    # 2. Actividad Sectorial (Página 1)
    h2 = doc.add_heading("2. Actividad Económica y Desagregación Sectorial en Cuyo", level=2)
    h2.paragraph_format.space_before = Pt(4)
    h2.paragraph_format.space_after = Pt(4)
    p2 = doc.add_paragraph(
        "El Estimador Mensual de Actividad Económica (EMAE) consolidó un avance de 3,1% interanual. En el ámbito regional, los despachos de vino al mercado interno informados por el Instituto Nacional de Vitivinicultura (INV) totalizaron 780 miles de hectolitros, marcando una estabilización en los niveles de consumo y comercialización de la cadena agroindustrial cuyana."
    )
    p2.paragraph_format.space_after = Pt(6)
    
    # Salto de página explícito para página 2
    doc.add_page_break()
    
    # 3. Balance del BCRA (Página 2)
    h3 = doc.add_heading("3. Saneamiento Cuasifiscal, Base Monetaria y Reservas Internacionales", level=2)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(4)
    p3 = doc.add_paragraph(
        "La autoridad monetaria completó la extinción de los pasivos remunerados de corto plazo (Pases pasivos), eliminando la fuente endógena de emisión por intereses cuasifiscales. El stock de Letras Fiscales de Liquidez (Lefi) emitidas por el Tesoro totalizó $33,5 billones, mientras que las reservas internacionales brutas se situaron en USD 27.129 millones con reservas netas positivas bajo métrica FMI."
    )
    p3.paragraph_format.space_after = Pt(6)
    
    # Figura Monetaria (Página 2)
    fig2_path = os.path.join(dir_figuras, "Dinamica_Monetaria_BCRA_v2.png")
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(6.3))
        p_cap2 = doc.add_paragraph()
        run_cap2 = p_cap2.add_run("Figura 1. Saneamiento del balance del BCRA, absorción por Lefi y reservas internacionales.")
        run_cap2.font.size = Pt(8.0)
        run_cap2.font.bold = True
        p_cap2.paragraph_format.space_after = Pt(10)
        
    # 4. Conclusiones Regionales (Página 2)
    h4 = doc.add_heading("4. Conclusiones y Perspectivas Macroeconómicas", level=2)
    h4.paragraph_format.space_before = Pt(4)
    h4.paragraph_format.space_after = Pt(4)
    p4 = doc.add_paragraph(
        "El régimen macroeconómico presenta señales claras de consolidación en el ancla nominal y fiscal. Para la región de Cuyo, la combinación de estabilidad cambiaria y reducción de la inflación general permite proyectar una recomposición gradual del poder de compra salarial, mitigando los costos de financiamiento para el sector productivo vitivinícola y comercial."
    )
    p4.paragraph_format.space_after = Pt(4)

    os.makedirs(os.path.dirname(os.path.abspath(ruta_salida_docx)), exist_ok=True)
    doc.save(ruta_salida_docx)
    return ruta_salida_docx


def exportar_lote_docx_a_pdf(pares_docx_pdf):
    import pythoncom
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        for ruta_docx, ruta_pdf in pares_docx_pdf:
            if os.path.exists(ruta_docx):
                doc_w = word.Documents.Open(os.path.abspath(ruta_docx), ReadOnly=True)
                doc_w.SaveAs(os.path.abspath(ruta_pdf), FileFormat=17) # 17 = wdFormatPDF
                doc_w.Close(SaveChanges=False)
                print(f"Exportado exitosamente: {os.path.basename(ruta_pdf)}")
    finally:
        word.Quit()
        pythoncom.CoUninitialize()


def exportar_docx_a_pdf(ruta_docx: str, ruta_pdf: str) -> str:
    exportar_lote_docx_a_pdf([(ruta_docx, ruta_pdf)])
    return ruta_pdf


if __name__ == "__main__":
    base = os.path.dirname(os.path.dirname(__file__))
    excel = os.path.join(base, "01_Bases_Datos", "Base_Datos_Macro_Financiera.xlsx")
    figuras = os.path.join(base, "03_Figuras_HD")
    
    docx_sem = os.path.join(base, "04_Informes_Semanales_APA7", "2026-08-21_Paper_Macroeconomico_Semanal.docx")
    pdf_sem = docx_sem.replace(".docx", ".pdf")
    compilar_paper_semanal(excel, figuras, docx_sem)
    
    docx_oeru = os.path.join(base, "05_Informes_Mensuales_OERU", "2026-08_Informe_Mensual_Coyuntura_OERU.docx")
    pdf_oeru = docx_oeru.replace(".docx", ".pdf")
    compilar_informe_oeru(excel, figuras, docx_oeru)
    
    exportar_lote_docx_a_pdf([
        (docx_sem, pdf_sem),
        (docx_oeru, pdf_oeru)
    ])
    print("Todos los informes fueron compilados a DOCX y PDF.")
