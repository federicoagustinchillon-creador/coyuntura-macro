"""
GENERADOR DEL MONITOR DIARIO DE MERCADOS Y COYUNTURA FINANCIERA (2 PÁGINAS TRADING DESK)
========================================================================================
Autor: Federico Agustín Chillón
Facultad de Ciencias Económicas — UNCUYO / OERU
Genera el monitor diario en 2 páginas institucionales con tipografía Georgia 9.2 pt,
paleta Oxford Navy / Deep Wine, análisis cuantitativo de cierre de mercados, flujos MULC,
arbitrajes de tasas y trade ideas ejecutables.

CORRECCIÓN DE ESTA SESIÓN: la auditoría del archivo encontró que este generador nunca
cargaba 01_Bases_Datos/datos_del_dia.json -- todo el contenido numérico (dólares, EMBI+,
TIRes soberanas, tasas de Lecap, múltiplos de acciones, licitaciones del Tesoro) era un
literal de plantilla fijo, incluyendo dos inconsistencias reales frente al contrato:
breakeven inflacionario mostrado en 2,65% cuando el dato real es 2,86%, y TIR de GD30
mostrada en 10,70% cuando el dato real es 9,80%. Se reemplaza cada bloque hardcodeado por
lectura dinámica de src/contexto_informe.py (punto único de datos reales) y de los
fetchers ya construidos (src/fetch_datos_reales.py, src/fetch_series_indec_bcra.py). Todo
campo sin fuente automatizable en el repo queda explícitamente marcado "s/d (sin fuente
automatizable)" en el propio texto -- nunca se omite en silencio ni se deja un número fijo
sin decir nada.
"""

import os
import sys
from datetime import datetime

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import win32com.client
import pandas as pd

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

from src.contexto_informe import cargar_contexto, fmt_pct, fmt_num
from src.fetch_datos_reales import obtener_variacion_semanal_acciones
from src.fetch_series_indec_bcra import obtener_monetario_reciente

COLOR_NAVY = RGBColor(12, 35, 64)       # Oxford Navy #0C2340
COLOR_WINE = RGBColor(114, 47, 55)      # Deep Wine #722F37
COLOR_CHARCOAL = RGBColor(15, 23, 42)   # Slate Charcoal #0F172A
COLOR_MUTED = RGBColor(100, 116, 139)   # Slate Gray #64748B
COLOR_FOREST = RGBColor(13, 92, 70)     # Forest Green #0D5C46
COLOR_OCHRE = RGBColor(180, 83, 9)      # Warm Amber #B45309

SIN_FUENTE = "s/d, sin fuente automatizable"

_MESES_ES = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio",
             "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]


def _fecha_larga(fecha_iso):
    """Convierte 'YYYY-MM-DD' (campo `fecha` del contrato) a 'D de Mes de AAAA'.
    Si no hay fecha cargada en el contrato, usa la fecha real de ejecución
    (datetime.now()) -- nunca la fecha de plantilla hardcodeada ("21 de Agosto
    de 2026") que tenía este archivo antes de esta corrección."""
    dt = None
    if fecha_iso:
        try:
            dt = datetime.strptime(fecha_iso, "%Y-%m-%d")
        except (ValueError, TypeError):
            dt = None
    if dt is None:
        dt = datetime.now()
    return f"{dt.day} de {_MESES_ES[dt.month]} de {dt.year}"


def _brecha_pct(nivel, base):
    """Brecha porcentual real entre dos niveles de tipo de cambio del propio
    contrato (ej. MEP vs. oficial minorista) -- una división, no un número
    inventado. Devuelve None si falta cualquiera de los dos insumos."""
    if nivel is None or base is None or base == 0:
        return None
    return (nivel / base - 1) * 100


def _tna_tem_str(tem):
    """El contrato (tasas_ars.*) solo publica la Tasa Efectiva Mensual (TEM)
    de Lecap. La Tasa Nominal Anual (TNA) que se muestra junto a ella es una
    anualización simple (TEM x 12), la misma convención que ya usaba este
    informe y src/generador_graficos_hd.py -- no es un dato adicional de
    fuente distinta, es la misma TEM real expresada en otra base."""
    if tem is None:
        return "s/d"
    tna = tem * 12
    return f"{fmt_pct(tna, 2)} TNA ({fmt_pct(tem, 2)} TEM)"


def _boncer_str(tir_real):
    if tir_real is None:
        return "s/d"
    return f"CER + {fmt_pct(tir_real, 2)} TIR Real"


def _tesis_desde_bl(black_litterman_views, asset_ticker):
    """Busca una vista táctica REAL cargada en black_litterman_tactical_views
    del contrato para el ticker dado. Si no hay vista cargada, se declara
    'Neutral (sin view táctico cargado)' en vez de inventar una recomendación
    de research que el contrato no respalda."""
    for view in black_litterman_views or []:
        if view.get("asset") == asset_ticker:
            thesis = view.get("thesis", "")
            retorno = view.get("view_return_annual", view.get("excess_return_annual"))
            etiqueta = "Sobreponderar" if (retorno or 0) > 0 else "Subponderar"
            return f"{etiqueta} · {thesis}" if thesis else etiqueta
    return "Neutral (sin view táctico cargado)"


def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top="CBD5E1", bottom="CBD5E1", left=None, right=None, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b_name, b_color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if b_color:
            b_el = OxmlElement(f'w:{b_name}')
            b_el.set(qn('w:val'), 'single')
            b_el.set(qn('w:sz'), str(sz))
            b_el.set(qn('w:space'), '0')
            b_el.set(qn('w:color'), b_color)
            tcBorders.append(b_el)
        else:
            b_el = OxmlElement(f'w:{b_name}')
            b_el.set(qn('w:val'), 'none')
            tcBorders.append(b_el)
    tcPr.append(tcBorders)

def add_header_footer_diario(doc, fecha_str=None):
    if fecha_str is None:
        fecha_str = _fecha_larga(None)
    for i, section in enumerate(doc.sections):
        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.text = f"MONITOR DIARIO DE MERCADOS & COYUNTURA · {fecha_str.upper()}\t\tFEDERICO AGUSTÍN CHILLÓN"
        p_hdr.runs[0].font.name = "Georgia"
        p_hdr.runs[0].font.size = Pt(7.8)
        p_hdr.runs[0].font.color.rgb = COLOR_MUTED

        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.text = "Federico Agustín Chillón · Investigador · Cs. Económicas UNCUYO · FCE UNCUYO\t\tMonitor Flash Diario"
        p_ftr.runs[0].font.name = "Georgia"
        p_ftr.runs[0].font.size = Pt(7.8)
        p_ftr.runs[0].font.color.rgb = COLOR_MUTED

def add_body(doc, text, space_after=3.5, font_size=8.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(font_size)
    r.font.color.rgb = COLOR_CHARCOAL
    return p

def crear_cuadro_estrategia_desk(doc, titulo, tesis_str, detalles_str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.20)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=18, bottom=18, left=24, right=24)
    set_cell_borders(cell, top="0C2340", bottom="0C2340", left="0C2340", right="0C2340", sz="8")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    r_tit = p.add_run(titulo.upper() + "\n")
    r_tit.font.name = "Georgia"; r_tit.font.size = Pt(8.2); r_tit.font.bold = True; r_tit.font.color.rgb = COLOR_NAVY

    r_f = p.add_run(tesis_str + "\n")
    r_f.font.name = "Georgia"; r_f.font.size = Pt(8.0); r_f.font.bold = True; r_f.font.color.rgb = COLOR_WINE

    r_exp = p.add_run(detalles_str)
    r_exp.font.name = "Georgia"; r_exp.font.size = Pt(7.8); r_exp.font.color.rgb = COLOR_CHARCOAL

def formatear_tabla_diaria(tabla, col_widths, headers, data_rows, font_size=7.4, alignments=None):
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tabla.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0C2340")
        set_cell_margins(hdr_cells[i], top=16, bottom=16, left=18, right=18)
        set_cell_borders(hdr_cells[i], top="0C2340", bottom="0C2340", left=None, right=None)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Georgia"; run.font.size = Pt(font_size); run.font.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row_data in enumerate(data_rows):
        row_cells = tabla.add_row().cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=12, bottom=12, left=18, right=18)
            set_cell_borders(row_cells[c_idx], top="E2E8F0", bottom="E2E8F0", left=None, right=None)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif ('%' in str(val) or '$' in str(val) or any(c.isdigit() for c in str(val))) and not ('Sobreponderar' in str(val) or 'Neutral' in str(val) or 'Subponderar' in str(val) or 'Lecap' in str(val) or 'Boncer' in str(val) or 'Oficial' in str(val) or 'MEP' in str(val) or 'CCL' in str(val) or 'Rofex' in str(val) or 'AL30' in str(val) or 'GD30' in str(val) or 'Tramos' in str(val) or 'Adjudicación' in str(val) or 'Letras' in str(val)):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in p.runs:
                run.font.name = "Georgia"; run.font.size = Pt(font_size)
                if 'Sobreponderar' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_FOREST
                elif 'Neutral' in str(val) or 'Mantener' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_OCHRE
                elif 'Subponderar' in str(val) or 'Reducir' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_WINE
                elif '+' in str(val) and '%' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_FOREST
                elif '-' in str(val) and '%' in str(val) and not ('USD' in str(val) or '2026' in str(val) or '$' in str(val)):
                    run.font.bold = True; run.font.color.rgb = COLOR_WINE
                else:
                    run.font.color.rgb = COLOR_CHARCOAL

    for row in tabla.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

def compilar_informe_diario(ruta_salida_docx: str, fecha_str: str = None) -> str:
    # --- Carga única de datos reales (src/contexto_informe.py) ------------
    ctx = cargar_contexto(incluir_series_lentas=False)
    dolar = ctx["dolar"]
    tasas_ars = ctx["tasas_ars"]
    soberano_usd = ctx["soberano_usd"]
    equity = ctx["equity"]
    tasas_bcra_ref = ctx["tasas_bcra_referencia"]
    black_litterman = ctx["black_litterman_tactical_views"]

    if fecha_str is None:
        fecha_str = _fecha_larga(ctx["datos"].get("fecha"))

    try:
        variacion_semanal = obtener_variacion_semanal_acciones()
    except Exception as e:
        print(f"      [Informe Diario] ERROR variación semanal de acciones: {e}")
        variacion_semanal = {}

    try:
        monetario = obtener_monetario_reciente()
    except Exception as e:
        print(f"      [Informe Diario] ERROR monetario BCRA: {e}")
        monetario = None

    al30_tir = soberano_usd.get("al30_tir")
    gd30_tir = soberano_usd.get("gd30_tir")
    gd35_tir = soberano_usd.get("gd35_tir")
    gd38_tir = soberano_usd.get("gd38_tir")
    embi = soberano_usd.get("embi_riesgo_pais_pbs")
    # Variacion diaria del EMBI+: fuente secundaria real (ArgentinaDatos,
    # agregador comunitario -- no JP Morgan/Bloomberg directo), solo para
    # el delta; el nivel sigue viniendo del contrato (fuente primaria).
    embi_var_1d = ctx.get("riesgo_pais_variacion_1d")
    embi_var_1d_pb = embi_var_1d["variacion_pb"] if embi_var_1d else None
    # Spread de legislación AL30/GD30: diferencia directa entre dos TIRes
    # reales del contrato, no un valor de plantilla ("50 pb" en la versión
    # anterior, que no coincidía con los 140 pb que arroja el dato real).
    spread_legislacion_pb = round((al30_tir - gd30_tir) * 100) if al30_tir is not None and gd30_tir is not None else None

    pases_ref = tasas_bcra_ref.get("pases_1d_tna", {})
    badlar_ref = tasas_bcra_ref.get("badlar_privados_tna", {})
    reservas_ref = tasas_bcra_ref.get("reservas_brutas_usd_m", {})

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.50)
        section.bottom_margin = Inches(0.50)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0.30)
        section.footer_distance = Inches(0.30)
        section.different_first_page_header_footer = True

    add_header_footer_diario(doc, fecha_str)

    # -------------------------------------------------------------------------
    # PÁGINA 1: TABLERO DE CIERRE, FLUJOS MULC & MICROESTRUCTURA CAMBIARIA
    # -------------------------------------------------------------------------
    p_title = doc.add_paragraph()
    r_t = p_title.add_run("Monitor Diario de Mercados & Coyuntura Financiera")
    r_t.font.name = "Georgia"; r_t.font.size = Pt(13.5); r_t.font.bold = True; r_t.font.color.rgb = COLOR_NAVY
    p_title.paragraph_format.space_before = Pt(0); p_title.paragraph_format.space_after = Pt(1.5)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run(f"Cierre de Operaciones: {fecha_str} | Estrategia Cuantitativa & Trading Desk | FCE UNCUYO")
    r_sub.font.name = "Georgia"; r_sub.font.size = Pt(8.0); r_sub.font.italic = True; r_sub.font.color.rgb = COLOR_MUTED
    p_sub.paragraph_format.space_after = Pt(3.5)

    # Tablero de 4 Indicadores Clave -- valores reales del contrato; toda
    # variación diaria que el contrato no releva queda marcada explícitamente
    # en vez de mostrarse como un porcentaje inventado.
    t_kpi = doc.add_table(rows=1, cols=4)
    t_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    merval_ars = equity.get("merval_ars")
    merval_usd_ccl = equity.get("merval_usd_ccl")
    kpis = [
        ("DÓLAR CCL (CABLE)", fmt_num(dolar.get("ccl"), 2, "$"),
         f"Var. diaria s/d · Brecha oficial {fmt_pct(dolar.get('brecha_ccl_oficial_pct'), 2, True)}"),
        ("RIESGO PAÍS EMBI+", (fmt_num(embi, 0) + " pb") if embi is not None else "s/d",
         (f"Var. 1D {embi_var_1d_pb:+d} pb (ArgentinaDatos)" if embi_var_1d_pb is not None else "Var. diaria s/d (sin fuente automatizable)")),
        ("TCR BILATERAL ARS/USD", (fmt_num(ctx["tcr_bilateral"]["ultimo"]["tcr_indice"], 1)) if ctx.get("tcr_bilateral") and ctx["tcr_bilateral"].get("ultimo") else "s/d",
         (f"Base {ctx['tcr_bilateral']['base_mes']}=100 · {ctx['tcr_bilateral']['ultimo']['mes']}") if ctx.get("tcr_bilateral") and ctx["tcr_bilateral"].get("ultimo") else "Sin cache (correr fetch_tcr_bilateral.py)"),
        ("S&P MERVAL", (fmt_num(merval_ars, 0) + " pts") if merval_ars is not None else "s/d",
         f"Var. semanal {fmt_pct(equity.get('var_semanal_pct'), 2, True)} · USD CCL {fmt_num(merval_usd_ccl, 2)}")
    ]
    for i, (k_t, k_v, k_s) in enumerate(kpis):
        cell = t_kpi.rows[0].cells[i]
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=12, bottom=12, left=16, right=16)
        set_cell_borders(cell, top="0C2340", bottom="CBD5E1", left="CBD5E1", right="CBD5E1", sz="6")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

        rt = p.add_run(k_t.upper() + "\n")
        rt.font.name = "Georgia"; rt.font.size = Pt(6.8); rt.font.bold = True; rt.font.color.rgb = COLOR_NAVY

        rv = p.add_run(k_v + "\n")
        rv.font.name = "Georgia"; rv.font.size = Pt(9.8); rv.font.bold = True; rv.font.color.rgb = COLOR_CHARCOAL

        rs = p.add_run(k_s)
        rs.font.name = "Georgia"; rs.font.size = Pt(6.6); rs.font.color.rgb = COLOR_MUTED

    for c in t_kpi.rows[0].cells: c.width = Inches(1.80)

    # 1. Microestructura Cambiaria
    h1 = doc.add_heading("1. Microestructura Cambiaria, Flujos en el MULC y Futuros Matba-Rofex", level=2)
    h1.paragraph_format.space_before = Pt(5); h1.paragraph_format.space_after = Pt(2)
    for r in h1.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY

    reservas_valor = reservas_ref.get("valor")
    reservas_fecha = reservas_ref.get("fecha", "s/d")
    add_body(doc, (
        f"La rueda cambiaria cerró sin volumen operado del Mercado Único y Libre de Cambios (MULC) disponible de forma "
        f"automatizada en este pipeline ({SIN_FUENTE}). El tipo de cambio mayorista de referencia (Com. \"A\" 3500) se ubicó "
        f"en {fmt_num(dolar.get('mayorista'), 2, '$')}, con variación diaria {SIN_FUENTE}. Las reservas brutas del Banco "
        f"Central se ubicaron en USD {fmt_num(reservas_valor, 0)} M (registro interno, {reservas_fecha}); no hay conector "
        f"automatizado a compras/ventas netas diarias del BCRA en el MULC ni a saldo comprador acumulado en el mes ({SIN_FUENTE})."
    ), space_after=2.5, font_size=8.2)
    add_body(doc, (
        f"En el segmento financiero, el Dólar MEP cerró en {fmt_num(dolar.get('mep'), 2, '$')} y el Dólar Contado con "
        f"Liquidación (CCL) en {fmt_num(dolar.get('ccl'), 2, '$')} (variación diaria {SIN_FUENTE} en ambos casos), con una "
        f"brecha de {fmt_pct(dolar.get('brecha_ccl_oficial_pct'), 2, True)} del CCL frente al dólar oficial minorista "
        f"({fmt_num(dolar.get('oficial_bna'), 2, '$')}). El dólar blue (paralelo informal) cotizó en "
        f"{fmt_num(dolar.get('blue'), 2, '$')}, cotización de referencia sin mercado formal ni volumen verificable."
    ), space_after=2.5, font_size=8.2)
    add_body(doc, (
        "En el mercado de futuros de Matba-Rofex, este pipeline no cuenta con un conector automatizado (feed pago, sin API "
        f"pública integrada al repositorio): interés abierto, concentración por vencimiento y tasas nominales anuales "
        f"implícitas quedan como {SIN_FUENTE} en esta corrida."
    ), space_after=3.0, font_size=8.2)

    oficial_bna = dolar.get("oficial_bna")
    mayorista = dolar.get("mayorista")
    mep = dolar.get("mep")
    ccl = dolar.get("ccl")
    blue = dolar.get("blue")

    # Var. Diaria y Volumen Operado no tienen fuente automatizable para
    # ningun segmento (sin conector a flujos del MULC ni a volumen de
    # mercado) -- se sacan ambas columnas en vez de mostrarlas en "s/d"
    # fila tras fila. La fila de Rofex se saca tambien: ya esta cubierta
    # arriba en texto y en el grafico de dolar futuro implicito por CIP.
    t_fx = doc.add_table(rows=1, cols=3)
    formatear_tabla_diaria(
        t_fx,
        col_widths=[2.30, 1.60, 1.90],
        headers=["Segmento / Activo", "Cierre Spot ($)", "Brecha vs. Oficial (%)"],
        data_rows=[
            ["Dólar Oficial Minorista (BNA)", fmt_num(oficial_bna, 2, "$"), "0,00% (Base)"],
            ["Dólar Mayorista (A3500)", fmt_num(mayorista, 2, "$"), fmt_pct(_brecha_pct(mayorista, oficial_bna), 2, True)],
            ["Dólar MEP (Bolsa)", fmt_num(mep, 2, "$"), fmt_pct(_brecha_pct(mep, oficial_bna), 2, True)],
            ["Dólar CCL (Cable)", fmt_num(ccl, 2, "$"), fmt_pct(dolar.get("brecha_ccl_oficial_pct"), 2, True)],
            ["Dólar Blue (Informal)", fmt_num(blue, 2, "$"), fmt_pct(_brecha_pct(blue, oficial_bna), 2, True)],
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=7.2
    )

    # img_p11_1_13.png era un resto huerfano de una version vieja del
    # pipeline (pre-matplotlib, datos desactualizados) -- se apunta al
    # chart real y ya corregido que genera src/generador_graficos_hd.py.
    fig_fx = os.path.join(BASE_DIR, "03_Figuras_HD", "chart_indec_6_fx.png")
    if os.path.exists(fig_fx):
        p_pic = doc.add_paragraph()
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pic.paragraph_format.space_before = Pt(3); p_pic.paragraph_format.space_after = Pt(0)
        p_pic.add_run().add_picture(fig_fx, width=Inches(7.20))

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 2: RENTA FIJA SOBERANA, DESK CALLS & RESUMEN MONETARIO
    # -------------------------------------------------------------------------
    h2 = doc.add_heading("2. Renta Fija en ARS (LECAPS/BONCER) y Curva Soberana en Hard Dollar", level=2)
    h2.paragraph_format.space_before = Pt(0); h2.paragraph_format.space_after = Pt(2)
    for r in h2.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY

    add_body(doc, (
        f"En el mercado en moneda local, las Letras del Tesoro (Lecap) operaron con tasas efectivas mensuales (TEM) de "
        f"{fmt_pct(tasas_ars.get('lecap_corta_tem'), 2)} en el tramo corto y {fmt_pct(tasas_ars.get('lecap_larga_tem'), 2)} "
        f"en el tramo largo -- el contrato no especifica tickers ni vencimientos puntuales (ej. \"S31O6\"), por lo que no "
        f"se identifica una letra específica -- ofreciendo una prima de {fmt_num(tasas_ars.get('premio_tasa_fija_pbs'), 0)} "
        f"pb mensuales sobre la inflación esperada REM ({fmt_pct(tasas_ars.get('inflacion_esperada_rem_tem'), 2)} TEM). El "
        f"título ajustable por CER Boncer TZX27 rindió {_boncer_str(tasas_ars.get('boncer_tzx27_tir_real'))} (el contrato "
        f"no releva otros tramos de la curva CER), convalidando un breakeven inflacionario implícito de "
        f"{fmt_pct(tasas_ars.get('breakeven_inflacion_tem'), 2)} mensual."
    ), space_after=2.5, font_size=8.2)
    add_body(doc, (
        f"En deuda hard dollar, el riesgo país (EMBI+) se ubicó en {(fmt_num(embi, 0) + ' pb') if embi is not None else 's/d'} "
        f"(variación 1D: {(f'{embi_var_1d_pb:+d} pb, fuente secundaria ArgentinaDatos' if embi_var_1d_pb is not None else SIN_FUENTE)}). "
        f"Los títulos Globales bajo ley extranjera cotizaron con TIR de "
        f"{fmt_pct(gd30_tir, 2)} en el GD30, {fmt_pct(gd35_tir, 2)} en el GD35 y {fmt_pct(gd38_tir, 2)} en el GD38 "
        f"(precios en pesos/USD y duration modificada: {SIN_FUENTE}). El spread de legislación frente al Bonar bajo ley "
        f"local AL30 ({fmt_pct(al30_tir, 2)} TIR) se ubicó en "
        f"{(fmt_num(spread_legislacion_pb, 0) + ' pb') if spread_legislacion_pb is not None else 's/d'}, calculado como "
        f"diferencia directa de TIR entre ambas especies."
    ), space_after=2.5, font_size=8.2)
    add_body(doc, (
        f"En el segmento monetario, la tasa de caución bursátil a 1 día en ByMA cotiza {SIN_FUENTE}; como referencia de "
        f"tasa corta del sistema, la tasa de pases pasivos del BCRA a 1 "
        f"día se ubicó en {fmt_pct(pases_ref.get('valor'), 2)} TNA y la tasa BADLAR de bancos privados en "
        f"{fmt_pct(badlar_ref.get('valor'), 2)} TNA (ambas de registro interno, {pases_ref.get('fecha', 's/d')})."
    ), space_after=2.5, font_size=8.2)
    add_body(doc, (
        f"En el crédito corporativo privado, este pipeline no cuenta con un conector a Obligaciones Negociables (ONs) hard "
        f"dollar de emisores energéticos de primera línea: los rendimientos por tramo de vencimiento quedan como "
        f"{SIN_FUENTE}."
    ), space_after=2.5, font_size=8.2)
    add_body(doc, (
        f"La curva forward soberana de tasas terminales cotiza {SIN_FUENTE} en este pipeline; "
        "el modelo Nelson-Siegel calibrado sobre la curva vigente (β0, β1, β2, τ, R²) está disponible en "
        "soberano_usd.nelson_siegel del contrato y se referencia en la infografía de curva soberana correspondiente."
    ), space_after=3.0, font_size=8.2)

    # Precio, Var. 1D y Duration no tienen fuente automatizable para
    # ninguna especie de esta tabla (el contrato solo releva TIR/TNA) -- se
    # sacan las 3 columnas enteras en vez de mostrarlas en "s/d" fila tras
    # fila: si no hay dato real, no se lista la columna.
    t_bon = doc.add_table(rows=1, cols=3)
    formatear_tabla_diaria(
        t_bon,
        col_widths=[1.90, 1.60, 2.70],
        headers=["Ticker / Especie", "TIR / TNA (%)", "Tesis Desk"],
        data_rows=[
            ["Lecap tramo corto", _tna_tem_str(tasas_ars.get("lecap_corta_tem")), "Neutral (sin view táctico cargado)"],
            ["Lecap tramo largo", _tna_tem_str(tasas_ars.get("lecap_larga_tem")), "Neutral (sin view táctico cargado)"],
            ["Boncer TZX27", _boncer_str(tasas_ars.get("boncer_tzx27_tir_real")), "Neutral (sin view táctico cargado)"],
            ["Bonar 2030 (AL30)", (fmt_pct(al30_tir, 2) + " TIR") if al30_tir is not None else "s/d", "Neutral (sin view táctico cargado)"],
            ["Global 2030 (GD30)", (fmt_pct(gd30_tir, 2) + " TIR") if gd30_tir is not None else "s/d", _tesis_desde_bl(black_litterman, "GD30.BA")],
            ["Global 2035 (GD35)", (fmt_pct(gd35_tir, 2) + " TIR") if gd35_tir is not None else "s/d", "Neutral (sin view táctico cargado)"],
            ["Global 2038 (GD38)", (fmt_pct(gd38_tir, 2) + " TIR") if gd38_tir is not None else "s/d", "Neutral (sin view táctico cargado)"],
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=7.2
    )
    add_body(doc, "Precio, variación diaria y duration modificada de estas especies no tienen fuente automatizable en este pipeline (no hay motor de pricing de bonos en pesos conectado) -- se omiten en vez de mostrarse sin dato.",
             font_size=7.2, space_after=3.0)

    # Cuadro de Estrategia Táctica Cuantitativa Desk Call -- juicio de research
    # de mesa, no un dato de mercado; pero los insumos (TEM Lecap, TIR
    # soberanas, premio sobre REM) se leen del contrato real.
    crear_cuadro_estrategia_desk(
        doc,
        "Trade Call Cuantitativo de Mesa & Arbitraje de Spreads",
        f"Estrategia 1: Carry en Lecap tramo corto (TEM {fmt_pct(tasas_ars.get('lecap_corta_tem'), 2)}) | "
        f"Estrategia 2: Arbitraje AL30 vs. GD30 por legislación",
        (
            f"• Carry en Pesos: La TEM de {fmt_pct(tasas_ars.get('lecap_corta_tem'), 2)} del tramo corto de Lecap implica "
            f"una prima de {fmt_num(tasas_ars.get('premio_tasa_fija_pbs'), 0)} pb mensuales sobre la inflación esperada "
            f"REM ({fmt_pct(tasas_ars.get('inflacion_esperada_rem_tem'), 2)} TEM). El fondeo con caución a 1 día no tiene "
            f"tasa disponible en este pipeline: cotiza {SIN_FUENTE}; el spread de fondeo neto no puede calcularse de "
            f"forma automatizada en esta corrida.\n"
            f"• Arbitraje de Legislación: El diferencial de "
            f"{(fmt_num(spread_legislacion_pb, 0) + ' pb') if spread_legislacion_pb is not None else 's/d'} entre AL30 "
            f"({fmt_pct(al30_tir, 2)}) y GD30 ({fmt_pct(gd30_tir, 2)}) presenta una oportunidad de rotación táctica hacia "
            f"el título de mayor TIR, sujeta a riesgo de legislación y liquidez relativa.\n"
            f"• View Táctico Cargado en el Contrato: {_tesis_desde_bl(black_litterman, 'GD30.BA')}"
        )
    )

    # Monitor de Acciones Líderes y Cedears -- precio de cierre y variación
    # SEMANAL reales via yfinance (src/fetch_datos_reales.obtener_variacion_
    # semanal_acciones, que trae retornos semana a semana, no diarios -- la
    # columna se rotula en consecuencia). Múltiplos EV/EBITDA de equity.lideres
    # del contrato; view táctico de GGAL.BA leído de black_litterman_tactical_views.
    h3_eq = doc.add_heading("3. Monitor de Renta Variable Energética y Múltiplos Corporativos", level=2)
    h3_eq.paragraph_format.space_before = Pt(3); h3_eq.paragraph_format.space_after = Pt(2)
    for r in h3_eq.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY

    lideres_por_ticker = {l.get("ticker"): l for l in equity.get("lideres", [])}
    nombres_eq = {
        "YPFD": "YPF S.A. (YPFD)",
        "PAMP": "Pampa Energía (PAMP)",
        "TGSU2": "Transportadora Gas Sur (TGSU2)",
        "GGAL": "Grupo Financiero Galicia (GGAL)",
        "BMA": "Banco Macro (BMA)",
    }
    data_rows_eq = []
    for tk, nombre in nombres_eq.items():
        v_sem = variacion_semanal.get(tk)
        precio = fmt_num(v_sem["cierre_ars"], 2, "$") if v_sem else "s/d"
        var_sem = fmt_pct(v_sem["var_semanal_pct"], 2, True) if v_sem else "s/d"
        lid = lideres_por_ticker.get(tk)
        ev_ebitda = f"{fmt_num(lid.get('ev_ebitda'), 1)}x" if lid and lid.get("ev_ebitda") is not None else "s/d"
        if tk == "GGAL":
            tesis = _tesis_desde_bl(black_litterman, "GGAL.BA")
        elif lid and lid.get("recom"):
            tesis = lid["recom"].capitalize()
        else:
            tesis = "s/d (sin múltiplo/view cargado en el contrato)"
        data_rows_eq.append([nombre, precio, var_sem, ev_ebitda, tesis])

    t_eq_d = doc.add_table(rows=1, cols=5)
    formatear_tabla_diaria(
        t_eq_d,
        col_widths=[1.80, 1.20, 1.20, 1.30, 1.70],
        headers=["Compañía / Ticker", "Precio Cierre ($)", "Var. Semanal (%)", "EV/EBITDA", "Tesis Fundamental"],
        data_rows=data_rows_eq,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=7.0
    )

    # 4. Licitaciones del Tesoro y Drivers de la Próxima Rueda -- el contrato no
    # tiene conector a licitaciones primarias del Tesoro (montos adjudicados,
    # tasa de corte): la columna de tasa se rotula "mercado secundario" y usa
    # las TEM/TNA reales de tasas_ars, sin presentarlas como resultado de
    # licitación. LEFI se marca como mecanismo discontinuado (hallazgo del
    # commit 1b0a196: stock $0 desde jul-2025, BCRA id=196) en vez de repetir
    # el monto de plantilla ("$29,3 Billones") que contradice el dato real.
    # Base Monetaria y Pases Pasivos usan la serie real de BCRA (src/fetch_
    # series_indec_bcra.obtener_monetario_reciente), rotulados como promedio
    # mensual -- no como stock puntual a la fecha, que la serie no provee.
    h4 = doc.add_heading("4. Licitaciones del Tesoro, Absorción Monetaria y Drivers de la Próxima Rueda", level=2)
    h4.paragraph_format.space_before = Pt(3); h4.paragraph_format.space_after = Pt(2)
    for r in h4.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY

    base_m_str = "s/d"
    pases_m_str = "s/d"
    if monetario and monetario.get("base_m"):
        base_m_str = fmt_num(monetario["base_m"][-1], 1, "$") + " B"
    if monetario and monetario.get("pases_m"):
        pases_m_str = fmt_num(monetario["pases_m"][-1], 1, "$") + " B"
    pases_tna = pases_ref.get("valor")

    t_lic = doc.add_table(rows=1, cols=4)
    formatear_tabla_diaria(
        t_lic,
        col_widths=[2.10, 1.50, 1.50, 2.10],
        headers=["Instrumento / Operación", "Monto Adjudicado / Stock", "Tasa de Referencia (2rio.)", "Destino / Efecto Monetario"],
        data_rows=[
            ["Lecap tramo corto (mercado 2rio.)", "s/d (sin licitación primaria)", _tna_tem_str(tasas_ars.get("lecap_corta_tem")), "Referencia de tasa fija corta; sin dato de licitación."],
            ["Lecap tramo largo (mercado 2rio.)", "s/d (sin licitación primaria)", _tna_tem_str(tasas_ars.get("lecap_larga_tem")), "Referencia de tasa fija larga; sin dato de licitación."],
            ["Letras Fiscales de Liquidez (LEFI)", "Mecanismo discontinuado (stock $0 desde jul-2025, BCRA id=196)", "-", "Sin efecto monetario vigente."],
            ["Base Monetaria (promedio mensual)", base_m_str, "-", "BCRA v4.0, id=15; promedio mensual, no stock puntual a la fecha."],
            ["Pases Pasivos BCRA (1D)", pases_m_str, (fmt_pct(pases_tna, 2) + " TNA") if pases_tna is not None else "s/d", f"Absorción monetaria de corto plazo (registro interno, {pases_ref.get('fecha', 's/d')})."],
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=7.0
    )

    add_body(doc, "• Drivers para la Rueda Siguiente: (1) Licitación quincenal de títulos en pesos de la Secretaría de Finanzas; (2) Publicación de balances corporativos del sector energético (YPF, PAMP); (3) Evolución del saldo comprador en el MULC; (4) Vencimiento de futuros Rofex y volumen en el tramo corto.", font_size=8.0, space_after=0)

    os.makedirs(os.path.dirname(os.path.abspath(ruta_salida_docx)), exist_ok=True)
    doc.save(ruta_salida_docx)
    print("Monitor Diario DOCX generado:", ruta_salida_docx)
    return ruta_salida_docx

def compilar_y_exportar_informe_diario(ruta_docx: str, ruta_pdf: str, fecha_str: str = None):
    compilar_informe_diario(ruta_docx, fecha_str)
    import pythoncom
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc_w = word.Documents.Open(os.path.abspath(ruta_docx), ReadOnly=True)
        doc_w.SaveAs(os.path.abspath(ruta_pdf), FileFormat=17)
        doc_w.Close(SaveChanges=False)
        print(f"Monitor Diario PDF exportado exitosamente: {ruta_pdf}")
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

if __name__ == "__main__":
    _ctx_main = cargar_contexto(incluir_series_lentas=False)
    _fecha_iso = _ctx_main["datos"].get("fecha") or datetime.now().strftime("%Y-%m-%d")
    d_docx = os.path.join(BASE_DIR, "04_Informes_Diarios", f"{_fecha_iso}_Monitor_Diario_Mercados.docx")
    d_pdf = d_docx.replace(".docx", ".pdf")
    compilar_y_exportar_informe_diario(d_docx, d_pdf)
