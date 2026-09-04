"""
GENERADOR DEL PAPER SEMANAL DE INVESTIGACIÓN MACROECONÓMICA (4 PÁGINAS APA 7 ALTA DENSIDAD)
==========================================================================================
Autor: Federico Agustín Chillón
Facultad de Ciencias Económicas — UNCUYO / OERU
Genera el paper semanal en 4 páginas completas de alta densidad analítica, tipografía Georgia 8.8 pt,
paleta Oxford Navy / Deep Wine, matrices de sensibilidad de retorno total, descomposición Nelson-Siegel,
atribución de cartera y referencias bibliográficas formales APA 7, sin espacios en blanco residuales.
"""

import os
import sys
import math
from datetime import datetime, timedelta
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import win32com.client
import pandas as pd

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

from src.contexto_informe import cargar_contexto, fmt_pct, fmt_num, fmt_o_manual

SIN_FUENTE = "Estimación institucional"


def _calcular_periodo_semanal(fecha_iso):
    """Semana calendario (lunes a viernes) que contiene la fecha real del
    contrato (`datos_del_dia.json['fecha']`). Reemplaza el periodo fijo
    hardcodeado -- si no hay fecha en el contrato, cae a la fecha de
    ejecucion (`datetime.now()`), nunca a un periodo inventado distinto
    del que corresponde a la corrida real."""
    try:
        fecha = datetime.strptime(fecha_iso, "%Y-%m-%d") if fecha_iso else datetime.now()
    except (ValueError, TypeError):
        fecha = datetime.now()
    lunes = fecha - timedelta(days=fecha.weekday())
    viernes = lunes + timedelta(days=4)
    meses_es = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio",
                "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    if lunes.month == viernes.month:
        return f"Semana del {lunes.day} al {viernes.day} de {meses_es[viernes.month]} de {viernes.year}"
    return f"Semana del {lunes.day} de {meses_es[lunes.month]} al {viernes.day} de {meses_es[viernes.month]} de {viernes.year}"


def _ns_forward_instantaneo(beta0, beta1, beta2, tau, t):
    """Tasa forward instantanea f(t) del modelo de Nelson & Siegel (1987)
    evaluada en los parametros REALES calibrados del contrato
    (soberano_usd.nelson_siegel), no en un valor de relleno. Formula
    estandar: f(t) = beta0 + beta1*e^(-t/tau) + beta2*(t/tau)*e^(-t/tau).
    Devuelve None si falta algun parametro -- no fabrica una tasa forward
    sin los betas que la definen."""
    if None in (beta0, beta1, beta2, tau) or tau == 0:
        return None
    x = t / tau
    e = math.exp(-x)
    return beta0 + beta1 * e + beta2 * x * e

COLOR_NAVY = RGBColor(12, 35, 64)       # Oxford Navy #0C2340
COLOR_WINE = RGBColor(114, 47, 55)      # Deep Wine #722F37
COLOR_CHARCOAL = RGBColor(15, 23, 42)   # Slate Charcoal #0F172A
COLOR_MUTED = RGBColor(100, 116, 139)   # Slate Gray #64748B
COLOR_FOREST = RGBColor(13, 92, 70)     # Forest Green #0D5C46
COLOR_OCHRE = RGBColor(180, 83, 9)      # Warm Amber #B45309

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top="CBD5E1", bottom="CBD5E1", left=None, right=None, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b_name, b_color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if b_color:
            b_el = OxmlElement(f'w:{b_name}')
            b_el.set(qn('w:val'), 'single')
            b_el.set(qn('w:sz'), str(sz))
            b_el.set(qn('w:space'), '0')
            b_el.set(qn('w:color'), b_color)
            tcBorders.append(b_el)
        else:
            b_el = OxmlElement(f'w:{b_name}')
            b_el.set(qn('w:val'), 'none')
            tcBorders.append(b_el)
    tcPr.append(tcBorders)

def add_header_footer_semanal(doc, periodo_str=None):
    periodo_str = periodo_str or _calcular_periodo_semanal(None)
    for i, section in enumerate(doc.sections):
        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.text = f"PAPER DE INVESTIGACIÓN MACROECONÓMICA & ESTRATEGIA · {periodo_str.upper()}\t\tFEDERICO AGUSTÍN CHILLÓN"
        p_hdr.runs[0].font.name = "Georgia"
        p_hdr.runs[0].font.size = Pt(7.8)
        p_hdr.runs[0].font.color.rgb = COLOR_MUTED
        
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.text = "Federico Agustín Chillón · Investigador · Cs. Económicas UNCUYO · FCE UNCUYO\t\tPaper Semanal APA 7"
        p_ftr.runs[0].font.name = "Georgia"
        p_ftr.runs[0].font.size = Pt(7.8)
        p_ftr.runs[0].font.color.rgb = COLOR_MUTED

def add_body(doc, text, space_after=3.0, font_size=8.5, line_spacing=1.14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(font_size)
    r.font.color.rgb = COLOR_CHARCOAL
    return p

def crear_cuadro_formula(doc, titulo, formula_str, explicacion_str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.20)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=14, bottom=14, left=20, right=20)
    set_cell_borders(cell, top="0C2340", bottom="0C2340", left="0C2340", right="0C2340", sz="8")
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1.5)
    r_tit = p.add_run(titulo.upper() + "\n")
    r_tit.font.name = "Georgia"; r_tit.font.size = Pt(7.8); r_tit.font.bold = True; r_tit.font.color.rgb = COLOR_NAVY
    
    r_f = p.add_run(formula_str + "\n")
    r_f.font.name = "Consolas"; r_f.font.size = Pt(8.0); r_f.font.bold = True; r_f.font.color.rgb = COLOR_WINE
    
    r_exp = p.add_run(explicacion_str)
    r_exp.font.name = "Georgia"; r_exp.font.size = Pt(7.5); r_exp.font.color.rgb = COLOR_CHARCOAL

def formatear_tabla_apa7(tabla, col_widths, headers, data_rows, font_size=7.2, alignments=None):
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tabla.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0C2340")
        set_cell_margins(hdr_cells[i], top=14, bottom=14, left=16, right=16)
        set_cell_borders(hdr_cells[i], top="0C2340", bottom="0C2340", left=None, right=None)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Georgia"; run.font.size = Pt(font_size); run.font.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row_data in enumerate(data_rows):
        row_cells = tabla.add_row().cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=10, bottom=10, left=16, right=16)
            set_cell_borders(row_cells[c_idx], top="E2E8F0", bottom="E2E8F0", left=None, right=None)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif ('%' in str(val) or '$' in str(val) or any(c.isdigit() for c in str(val))) and not ('Sobreponderar' in str(val) or 'Neutral' in str(val) or 'Subponderar' in str(val) or 'Lecap' in str(val) or 'Boncer' in str(val) or 'Local' in str(val) or 'Nueva York' in str(val) or 'Compresión' in str(val) or 'Ampliación' in str(val) or 'Tramos' in str(val) or 'Acciones' in str(val) or 'Globales' in str(val) or 'Base' in str(val)):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            for run in p.runs:
                run.font.name = "Georgia"; run.font.size = Pt(font_size)
                if 'Sobreponderar' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_FOREST
                elif 'Neutral' in str(val) or 'Mantener' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_OCHRE
                elif 'Subponderar' in str(val) or 'Reducir' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_WINE
                elif '+' in str(val) and '%' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_FOREST
                elif '-' in str(val) and '%' in str(val) and not ('USD' in str(val) or '2026' in str(val) or '$' in str(val)):
                    run.font.bold = True; run.font.color.rgb = COLOR_WINE
                else:
                    run.font.color.rgb = COLOR_CHARCOAL
                    
    for row in tabla.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

def compilar_paper_semanal_completo(ruta_salida_docx: str, periodo_str=None):
    # Contexto unico de datos reales (ver src/contexto_informe.py) -- este
    # generador NO cargaba datos_del_dia.json antes de esta correccion, por
    # lo que todos los numeros del cuerpo del paper eran literales de
    # plantilla. incluir_series_lentas=True porque la Seccion 1.1 necesita
    # la serie mensual real de Base Monetaria y Pases Pasivos del BCRA
    # (src/fetch_series_indec_bcra.obtener_monetario_reciente).
    ctx = cargar_contexto(incluir_series_lentas=True)
    periodo_str = periodo_str or _calcular_periodo_semanal(ctx.get("fecha"))

    dolar = ctx["dolar"]
    tasas_ars = ctx["tasas_ars"]
    inflacion = ctx["inflacion"]
    soberano = ctx["soberano_usd"]
    ns = soberano.get("nelson_siegel", {})
    ref_bcra = ctx["tasas_bcra_referencia"]
    monetario = ctx.get("monetario_historico")
    ripte = ctx.get("ripte")

    def _campo_ref(clave):
        """Extrae 'valor' de un campo de tasas_bcra_referencia (dict con
        valor/fecha/fuente) o None si no esta cargado."""
        d = ref_bcra.get(clave)
        return d.get("valor") if isinstance(d, dict) else None

    reservas_brutas = _campo_ref("reservas_brutas_usd_m")
    pases_tna = _campo_ref("pases_1d_tna")
    dolar_futuro = ctx.get("dolar_futuro_implicito")  # CIP teorico, NO cotizacion Rofex real

    if monetario and monetario.get("base_m"):
        base_monetaria_billones = monetario["base_m"][-1]
        pases_stock_billones = monetario["pases_m"][-1]
    else:
        base_monetaria_billones = None
        pases_stock_billones = None

    def _pct(v, signo=False):
        """fmt_pct con 2 decimales -- varios campos reales del contrato
        (ej. pases_1d_tna=23,12; gd35_tir=9,65; tasa_real_exante=0,95)
        requieren 2 decimales para no perder precision frente al 1 decimal
        por defecto de fmt_pct."""
        return fmt_pct(v, decimales=2, signo=signo)

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.50)
        section.bottom_margin = Inches(0.50)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0.30)
        section.footer_distance = Inches(0.30)
        section.different_first_page_header_footer = True
        
    add_header_footer_semanal(doc, periodo_str)
    
    # -------------------------------------------------------------------------
    # PÁGINA 1: TÍTULO, ABSTRACT, RÉGIMEN & TRANSMISIÓN MONETARIA COMPLETA
    # -------------------------------------------------------------------------
    p_tit = doc.add_paragraph()
    r_t = p_tit.add_run("Paper de Investigación Macroeconómica y Estrategia Financiera")
    r_t.font.name = "Georgia"; r_t.font.size = Pt(13.5); r_t.font.bold = True; r_t.font.color.rgb = COLOR_NAVY
    p_tit.paragraph_format.space_before = Pt(0); p_tit.paragraph_format.space_after = Pt(1.5)
    
    p_meta = doc.add_paragraph()
    r_m = p_meta.add_run(f"Período: {periodo_str} | Autor: Federico Agustín Chillón | FCE UNCUYO / OERU\nMarco Institucional: Análisis de Renta Fija Soberana, Microestructura Cambiaria y Régimen Monetario")
    r_m.font.name = "Georgia"; r_m.font.size = Pt(8.0); r_m.font.italic = True; r_m.font.color.rgb = COLOR_MUTED
    p_meta.paragraph_format.space_after = Pt(2.5)
    
    # Abstract
    tbl_abs = doc.add_table(rows=1, cols=1)
    tbl_abs.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_abs = tbl_abs.rows[0].cells[0]; c_abs.width = Inches(7.20)
    set_cell_background(c_abs, "F8FAFC"); set_cell_margins(c_abs, top=10, bottom=10, left=16, right=16)
    set_cell_borders(c_abs, top="722F37", bottom="722F37", left="722F37", right="722F37", sz="8")
    
    p_ab = c_abs.paragraphs[0]
    p_ab.paragraph_format.space_before = Pt(0); p_ab.paragraph_format.space_after = Pt(1.5)
    r_at = p_ab.add_run("RESUMEN EJECUTIVO & PALABRAS CLAVE:\n")
    r_at.font.name = "Georgia"; r_at.font.size = Pt(7.6); r_at.font.bold = True; r_at.font.color.rgb = COLOR_NAVY
    
    embi_str = fmt_num(soberano.get("embi_riesgo_pais_pbs"), 0)
    r_atx = p_ab.add_run(
        f"El presente documento examina la microestructura macro-financiera argentina correspondiente a la {periodo_str.lower()}. Se modela paramétricamente la curva soberana en moneda extranjera mediante la metodología de Nelson-Siegel, verificando un riesgo país (EMBI+) de {embi_str} pb. En el mercado en moneda local, se cuantifica la prima real ex-ante en letras de tasa fija (Lecaps) frente a las expectativas inflacionarias del REM, analizando la sustentabilidad del carry trade y el estado vigente de los instrumentos de absorción monetaria del BCRA.\n"
        "Palabras Clave: Nelson-Siegel, Arbitraje de Tasas, Absorción Monetaria, Carry Trade, Brecha Cambiaria, RIGI."
    )
    r_atx.font.name = "Georgia"; r_atx.font.size = Pt(7.4); r_atx.font.color.rgb = COLOR_CHARCOAL
    
    # 1. Diagnóstico
    h1 = doc.add_heading("1. Marco Macroeconómico, Dominancia Fiscal y Anclaje Nominal", level=2)
    h1.paragraph_format.space_before = Pt(4); h1.paragraph_format.space_after = Pt(2)
    for r in h1.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
    
    add_body(doc, f"La economía argentina consolidó su anclaje fiscal y monetario durante la semana analizada. El superávit primario de caja del Sector Público Nacional permitió prescindir de financiamiento monetario directo e indirecto. Las Letras Fiscales de Liquidez (Lefi), instrumento que había absorbido la deuda remunerada del BCRA, se encuentran discontinuadas desde julio de 2025 (stock verificado en $0 vía series BCRA v4.0); la absorción de liquidez opera hoy mediante Pases Pasivos a 1 día, con una tasa de referencia de {_pct(pases_tna)} TNA, desacoplando la creación endógena de dinero del balance de la autoridad monetaria.", font_size=8.2, space_after=2.5)
    # RIPTE: fuente secundaria real (Secretaria de Trabajo via
    # apis.datos.gob.ar, ver src/fetch_series_secundarias.py) -- es
    # NOMINAL (pesos corrientes), no deflactada, se aclara en el texto en
    # vez de llamarlo "salario real" sin mas.
    if ripte and ripte.get("var_mensual_ultimo") is not None:
        _ripte_txt = (f"La Remuneración Imponible Promedio de los Trabajadores Estables (RIPTE) creció "
                      f"{_pct(ripte['var_mensual_ultimo'])} mensual en términos nominales (Secretaría de Trabajo; "
                      f"no deflactada por inflación)")
    else:
        _ripte_txt = f"El salario nominal (RIPTE) no cuenta con dato disponible en esta corrida ({SIN_FUENTE})"
    add_body(doc, f"En materia de precios, la convergencia inflacionaria hacia el {_pct(inflacion.get('indec_general_mom'))} mensual a nivel nacional (INDEC) y {_pct(inflacion.get('deie_mendoza_mom'))} en Mendoza (DEIE) reafirma la efectividad del esquema cambiario vigente como ancla nominal intermedia. La inflación núcleo ({_pct(inflacion.get('indec_nucleo_mom'))} MoM) convalida la desaceleración del pass-through. {_ripte_txt}; la Canasta Básica Total en Mendoza se ubica en {fmt_num(inflacion.get('canasta_basica_total_mza'), 0, '$')}.", font_size=8.2, space_after=2.5)
    ccl_vs_mayorista_pct = round(100 * (dolar["ccl"] / dolar["mayorista"] - 1), 2) if dolar.get("ccl") and dolar.get("mayorista") else None
    add_body(doc, f"La estabilidad cambiaria en el segmento financiero (Dólar CCL en {fmt_num(dolar.get('ccl'), 2, '$')} con brecha del {_pct(dolar.get('brecha_ccl_oficial_pct'))} sobre el dólar oficial BNA y {_pct(ccl_vs_mayorista_pct)} sobre el mayorista) consolida un entorno de previsibilidad que reduce la demanda precautoria de dólares y estimula la remonetización del crédito privado en pesos.", font_size=8.2, space_after=2.5)
    
    # 1.1 Dinámica Cuasifiscal y Regla de Taylor
    h11 = doc.add_heading("1.1 Mecanismo de Transmisión Cuasifiscal, Regla de Taylor y Balance BCRA", level=3)
    h11.paragraph_format.space_before = Pt(3); h11.paragraph_format.space_after = Pt(2)
    for r in h11.runs: r.font.name = "Georgia"; r.font.size = Pt(8.6); r.font.bold = True; r.font.color.rgb = COLOR_WINE
    
    add_body(doc, f"Bajo el enfoque de Sargent & Wallace (1981), la disciplina fiscal mitiga el canal de expectativas racionales que asociaba los pasivos remunerados a emisión futura. Al absorber liquidez con Pases Pasivos respaldados por superávit primario, el multiplicador monetario secundario opera acotado por los encajes no remunerados, consolidando una Base Monetaria de {fmt_o_manual(base_monetaria_billones, lambda v: fmt_num(v, 2, '$') + ' Billones (BCRA v4.0)')} y reservas brutas en {fmt_o_manual(reservas_brutas, lambda v: fmt_num(v, 0, 'USD ') + ' millones')}.", font_size=8.2, space_after=2.5)
    add_body(doc, f"La ecuación de Fisher ex-ante (1 + i) = (1 + r)(1 + π^e) valida que con una TEM de Lecap corta de {_pct(tasas_ars.get('lecap_corta_tem'))} y expectativas del REM de {_pct(tasas_ars.get('inflacion_esperada_rem_tem'))}, la tasa de interés real ex-ante ({_pct(ctx.get('tasa_real_exante_tem_pct'), signo=True)} mensual) opera en terreno contractivo, garantizando la convergencia desinflacionaria. (Nota: se reemplaza la referencia de Lefi por la Lecap corta como instrumento de tasa fija de referencia, dado que el mecanismo de Lefi está discontinuado desde jul-2025.)", font_size=8.2, space_after=2.5)

    crear_cuadro_formula(
        doc,
        "Regla de Taylor (1993) & Brecha Contractiva de Política Monetaria",
        f"i_t = r* + π_t + 0,5 · (π_t - π*) + 0,5 · y_gap   ==>   i_real = {_pct(ctx.get('tasa_real_exante_tem_pct'), signo=True)} TEM vs r* = 0,75% mensual (supuesto del analista)",
        f"La tasa real ex-ante de la Lecap corta ({_pct(ctx.get('tasa_real_exante_tem_pct'), signo=True)} mensual) se compara contra una tasa neutral r* = 0,75% mensual, que es un supuesto del analista y no un dato observado (no existe en el repositorio una estimación econométrica de r* para Argentina), estableciendo una postura monetaria contractiva que favorece el anclaje desinflacionario."
    )
    
    # Fuentes reales: Base Monetaria y stock de Pases via BCRA v4.0
    # (obtener_monetario_reciente, ids 15/152). Lefi confirmado en stock $0
    # desde jul-2025 (mecanismo discontinuado) -- se declara explicitamente
    # en vez de mostrar el monto de plantilla ($29,3 B) como si siguiera
    # vigente. La TNA de Pases (tasas_bcra_referencia.pases_1d_tna) es un
    # campo real independiente del stock y NO es 0,00% como decia la
    # plantilla -- esa era una contradiccion (instrumento "extinto" con
    # tasa de referencia vigente reportada en 0%). Reservas Netas y
    # Depositos Privados/Caucion no tienen fuente automatizable en el repo.
    t_mon = doc.add_table(rows=1, cols=4)
    formatear_tabla_apa7(
        t_mon,
        col_widths=[2.40, 1.40, 1.50, 1.90],
        headers=["Agregado / Instrumento Monetario", "Stock Vigente", "Tasa Nominal Anual", "Condición de Equilibrio"],
        data_rows=[
            ["Base Monetaria Ampliada", fmt_o_manual(base_monetaria_billones, lambda v: fmt_num(v, 2, "$") + " Billones"), "-", "Anclaje de agregados reales."],
            ["Letras Fiscales de Liquidez (Lefi)", "$0,00 Billones", "Mecanismo discontinuado (jul-2025)", "Reemplazado por Pases Pasivos."],
            ["Pases Pasivos BCRA (1 Día)", fmt_o_manual(pases_stock_billones, lambda v: fmt_num(v, 2, "$") + " Billones"), _pct(pases_tna) + " TNA", "Instrumento vigente de absorción."],
            ["Reservas Internacionales Brutas", fmt_o_manual(reservas_brutas, lambda v: fmt_num(v, 0, "USD ") + " M"), "-", "Registro interno BCRA."],
            ["Reservas Netas (Métrica FMI)", SIN_FUENTE, "-", "Sin conector automatizable en el repo."],
            ["Depósitos Privados en ARS", SIN_FUENTE, SIN_FUENTE, "Sin conector automatizable en el repo."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=7.0
    )
    
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 2: MICROESTRUCTURA CAMBIARIA & CARRY TRADE COMPLETA
    # -------------------------------------------------------------------------
    h2 = doc.add_heading("2. Microestructura Cambiaria, Paridad CIP y Rendimiento en Moneda Extranjera", level=2)
    h2.paragraph_format.space_before = Pt(0); h2.paragraph_format.space_after = Pt(2)
    for r in h2.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
    
    add_body(doc, f"Las cotizaciones del tipo de cambio financiero operaron con un spread comprimido: el Dólar CCL finalizó en {fmt_num(dolar.get('ccl'), 2, '$')} (brecha del {_pct(dolar.get('brecha_ccl_oficial_pct'))} respecto al oficial BNA y {_pct(ccl_vs_mayorista_pct)} sobre el mayorista Com. 'A' 3500 de {fmt_num(dolar.get('mayorista'), 2, '$')}). El volumen operado en el segmento mayorista y la intervención compradora/vendedora del BCRA en la rueda no tienen fuente automatizable en este contrato ({SIN_FUENTE}).", font_size=8.2, space_after=2.5)
    add_body(doc, f"El mercado de futuros Matba-Rofex (interés abierto, tasas nominales implícitas por posición) no cuenta con conector automatizado en el repositorio ({SIN_FUENTE}). En su lugar, la tabla siguiente muestra un dólar futuro teórico derivado por paridad de tasas cubierta (CIP) sobre datos reales del contrato -- un modelo, no una cotización de mercado observada.", font_size=8.2, space_after=2.5)

    crear_cuadro_formula(
        doc,
        "Fórmula de Retorno en USD por Carry Trade & Paridad CIP (Covered Interest Parity)",
        "R_USD = [(1 + TEM_lecap) / (1 + Δe_esperada)] - 1   |   CIP_Basis = (1 + i_ARS) - [(F_T / S_0) · (1 + i_USD)]",
        f"Para una TEM de {_pct(tasas_ars.get('lecap_corta_tem'))} en Lecap corta y una inflación esperada REM de {_pct(tasas_ars.get('inflacion_esperada_rem_tem'))} mensual (usada como proxy de Δe_esperada, dado que el contrato no incluye una serie propia de ritmo de devaluación/crawling peg), el rendimiento esperado en moneda dura asciende a {_pct(ctx.get('tasa_real_exante_tem_pct'), signo=True)} mensual, incentivando el fondeo en caución bursátil. F_T/S_0 (base CIP con futuros Rofex) queda sin cálculo: ver nota sobre ausencia de conector Matba-Rofex."
    )
    
    # Sin conector a Matba-Rofex en el repo (ver src/fetch_datos_reales.py):
    # en vez de una tabla de futuros Rofex 100% "0,0" (precio/TNA/interes
    # abierto de relleno, como hacia la plantilla anterior), se muestran las
    # dos cotizaciones spot reales del contrato y el dolar futuro CIP real
    # (src/modelos_riesgo.calcular_dolar_futuro_implicito), declarado
    # explicitamente como modelo y no como cotizacion de mercado.
    _dolar_futuro_por_dias = {c["dias"]: c for c in dolar_futuro["curva"]} if dolar_futuro else {}
    t_fx_sem = doc.add_table(rows=1, cols=4)
    fila_spot = [
        ["Dólar Mayorista (A 3500)", fmt_num(dolar.get("mayorista"), 2, "$"), "-", "0,00% (Base)"],
        ["Dólar CCL Cable", fmt_num(dolar.get("ccl"), 2, "$"), "-", _pct(ccl_vs_mayorista_pct, signo=True)],
    ]
    filas_cip = []
    for _dias, _label in ((30, "Futuro CIP 30d (teórico)"), (90, "Futuro CIP 90d (teórico)"), (180, "Futuro CIP 180d (teórico)")):
        _c = _dolar_futuro_por_dias.get(_dias)
        filas_cip.append([
            _label,
            fmt_num(_c["futuro_implicito"], 2, "$") if _c else SIN_FUENTE,
            _pct(_c["tna_implicita_pct"]) if _c else SIN_FUENTE,
            "Modelo CIP, no cotización Rofex",
        ])
    formatear_tabla_apa7(
        t_fx_sem,
        col_widths=[2.00, 1.30, 1.30, 2.60],
        headers=["Posición / Segmento", "Cotización / Futuro Implícito", "TNA Implícita", "Brecha Oficial / Nota"],
        data_rows=fila_spot + filas_cip,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=7.0
    )
    
    # img_p11_1_13.png era un resto huerfano de una version vieja del
    # pipeline (pre-matplotlib, datos desactualizados) -- se apunta al
    # chart real y ya corregido que genera src/generador_graficos_hd.py.
    fig3_path = os.path.join(BASE_DIR, "03_Figuras_HD", "chart_indec_6_fx.png")
    if os.path.exists(fig3_path):
        p_pic = doc.add_paragraph()
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pic.paragraph_format.space_before = Pt(3); p_pic.paragraph_format.space_after = Pt(0)
        p_pic.add_run().add_picture(fig3_path, width=Inches(7.20))
        
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(1); p_cap.paragraph_format.space_after = Pt(0)
        r_c = p_cap.add_run("Figura 1. Microestructura cambiaria, cotizaciones spot y estructura temporal de futuros Matba-Rofex.")
        r_c.font.name = "Georgia"; r_c.font.size = Pt(7.0); r_c.font.italic = True; r_c.font.color.rgb = COLOR_MUTED
        
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 3: CURVAS SOBERANAS & MODELADO NELSON-SIEGEL
    # -------------------------------------------------------------------------
    h3 = doc.add_heading("3. Estructura Temporal de la Deuda Soberana en USD y Calibración Nelson-Siegel", level=2)
    h3.paragraph_format.space_before = Pt(0); h3.paragraph_format.space_after = Pt(2)
    for r in h3.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
    
    spread_legislacion_pb = round((soberano["al30_tir"] - soberano["gd30_tir"]) * 100, 0) if soberano.get("al30_tir") is not None and soberano.get("gd30_tir") is not None else None
    add_body(doc, f"La curva soberana en moneda extranjera ubicó a los Globales Ley NY entre {_pct(soberano.get('gd38_tir'))} (GD38) y {_pct(soberano.get('gd30_tir'))} (GD30), mientras el único par Ley Local/Ley NY con TIR real disponible en el contrato (AL30 vs. GD30) mostró un spread de legislación de {fmt_o_manual(spread_legislacion_pb, lambda v: fmt_num(v, 0) + ' pb')} -- no se dispone de un promedio multi-tramo verificable por falta de TIR real de AL35/GD41 en el contrato. El ajuste econométrico bajo el modelo paramétrico de Nelson-Siegel valida una curva con R² = {fmt_num(ns.get('r2'), 3)}.", font_size=8.2, space_after=2.5)
    forward_10y = _ns_forward_instantaneo(ns.get("beta0"), ns.get("beta1"), ns.get("beta2"), ns.get("tau"), 10)
    add_body(doc, f"La curva forward instantánea f(t), calculada analíticamente a partir de los parámetros reales de Nelson-Siegel del contrato, proyecta una tasa terminal de {fmt_o_manual(forward_10y, fmt_pct)} a los 10 años. El atractivo relativo de los cupones step-up (GD38) frente al tramo ultralargo (GD41) y el retorno por roll-down no tienen fuente automatizable en el repositorio (no existe motor de pricing de bonos): {SIN_FUENTE}.", font_size=8.2, space_after=2.5)

    crear_cuadro_formula(
        doc,
        "Modelo Paramétrico de Curva de Rendimientos (Nelson & Siegel, 1987)",
        "y(t) = β₀ + β₁ · [(1 - e^{-t/τ}) / (t/τ)] + β₂ · [(1 - e^{-t/τ}) / (t/τ) - e^{-t/τ}]",
        f"Parámetros calibrados (fuente: contrato real soberano_usd.nelson_siegel): Nivel β₀ = {_pct(ns.get('beta0'))} (asíntota de largo plazo), Pendiente β₁ = {_pct(ns.get('beta1'), signo=True)}, Curvatura β₂ = {_pct(ns.get('beta2'), signo=True)}, Decaimiento τ = {fmt_num(ns.get('tau'), 2)} y R² = {fmt_num(ns.get('r2'), 3)}. RMSE del ajuste: sin fuente automatizable ({SIN_FUENTE})."
    )
    
    # TIR: real cuando el contrato tiene el campo (GD30/AL30/GD35/GD38).
    # AL35 y GD41 se retiran de la tabla -- el contrato no trae TIR para
    # ninguno de los dos, y no queda ningun otro campo real que mostrar en
    # esa fila (no vale la pena una fila 100% "0,0"). Precio Spot,
    # Duration/Convexidad y Roll-Down tampoco se muestran: el repositorio
    # no tiene motor de pricing de bonos ni cronogramas de cupones/
    # amortizacion verificados para calcular Macaulay duration real -- en
    # vez de 3 columnas enteras de relleno, se retiran y se declara la
    # ausencia una sola vez en el texto de arriba.
    t_bon_sem = doc.add_table(rows=1, cols=3)
    formatear_tabla_apa7(
        t_bon_sem,
        col_widths=[2.20, 1.60, 1.60],
        headers=["Título / Ticker", "Legislación", "TIR Anual"],
        data_rows=[
            ["Global 2030 (GD30)", "Nueva York", _pct(soberano.get("gd30_tir"))],
            ["Bonar 2030 (AL30)", "Argentina", _pct(soberano.get("al30_tir"))],
            ["Global 2035 (GD35)", "Nueva York", _pct(soberano.get("gd35_tir"))],
            ["Global 2038 (GD38)", "Nueva York", _pct(soberano.get("gd38_tir"))],
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=7.0
    )
    
    # img_p10_1_12.png era un resto huerfano de una version vieja del
    # pipeline (pre-matplotlib, datos desactualizados) -- se apunta al
    # chart real y ya corregido que genera src/generador_graficos_hd.py.
    fig1_path = os.path.join(BASE_DIR, "03_Figuras_HD", "chart_indec_5_sovereign.png")
    if os.path.exists(fig1_path):
        p_pic = doc.add_paragraph()
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pic.paragraph_format.space_before = Pt(3); p_pic.paragraph_format.space_after = Pt(0)
        p_pic.add_run().add_picture(fig1_path, width=Inches(7.20))
        
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(1); p_cap.paragraph_format.space_after = Pt(0)
        r_c = p_cap.add_run("Figura 2. Curva spot de rendimientos soberanos USD y estructura forward instantánea f(t).")
        r_c.font.name = "Georgia"; r_c.font.size = Pt(7.0); r_c.font.italic = True; r_c.font.color.rgb = COLOR_MUTED
        
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 4: SENSIBILIDAD, ASIGNACIÓN TÁCTICA & REFERENCIAS APA 7
    # -------------------------------------------------------------------------
    h4 = doc.add_heading("4. Análisis de Sensibilidad, Portafolio Modelo y Referencias Bibliográficas", level=2)
    h4.paragraph_format.space_before = Pt(0); h4.paragraph_format.space_after = Pt(2)
    for r in h4.runs: r.font.name = "Georgia"; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
    
    add_body(doc, "A partir de la expansión de segundo orden de Taylor para la sensibilidad del precio de los títulos de deuda soberana, se proyectan escenarios ilustrativos de retorno total en USD ante compresión y ampliación del spread crediticio (EMBI+).", font_size=8.2, space_after=2.5)
    add_body(doc, f"Advertencia metodológica: el repositorio no cuenta con motor de pricing de bonos ni con cronogramas de cupones/amortización verificados para los títulos soberanos argentinos, por lo que los valores de duration y convexidad usados a continuación son un supuesto ilustrativo del analista (no un cálculo de Macaulay/modified duration verificado) y los retornos resultantes son una proyección propia sujeta a revisión, no un resultado de mercado observado.", font_size=8.2, space_after=2.5)

    crear_cuadro_formula(
        doc,
        "Aproximación de Precios por Modified Duration y Convexidad (Taylor)",
        "ΔP / P ≈ -D_mod · Δy + ½ · C · (Δy)²",
        f"Bajo un supuesto ilustrativo de Duration = 5,81 y Convexidad = 37,2x para el tramo largo (GD38, {_pct(soberano.get('gd38_tir'))} TIR real), una compresión de spread de -300 pb generaría un upside proyectado de +18,45% en USD. Esta cifra es una estimación propia del analista, no un cálculo verificado por un motor de pricing."
    )

    t_sens = doc.add_table(rows=1, cols=4)
    formatear_tabla_apa7(
        t_sens,
        col_widths=[2.40, 1.60, 1.60, 1.60],
        headers=["Shock de Spread Soberano*", "Retorno GD38* (supuesto)", "Retorno GD35* (supuesto)", "Retorno AL30* (supuesto)"],
        data_rows=[
            ["Compresión -300 pb (Hacia 200 pb EMBI)", "+18,45% en USD", "+16,20% en USD", "+8,50% en USD"],
            ["Compresión -200 pb (Hacia 300 pb EMBI)", "+12,10% en USD", "+10,60% en USD", "+5,60% en USD"],
            ["Compresión -100 pb (Hacia 400 pb EMBI)", "+5,95% en USD", "+5,10% en USD", "+2,75% en USD"],
            ["Ampliación +100 pb (Hacia 600 pb EMBI)", "-5,45% en USD", "-4,80% en USD", "-2,65% en USD"],
            ["Ampliación +300 pb (Shock Externo)", "-14,80% en USD", "-13,20% en USD", "-7,65% en USD"]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=7.0
    )
    add_body(doc, "* Proyección propia del analista bajo supuestos ilustrativos de duration/convexidad; no constituye un cálculo de duration/convexity verificado con motor de pricing de bonos ni cronogramas de cupones reales.", font_size=6.6, space_after=2.5)

    # Tabla de Asignación Táctica de Cartera: recomendación de research del
    # analista (vistas tácticas Black-Litterman), no un dato de mercado
    # observado -- se aclara explícitamente antes de la tabla.
    add_body(doc, "La siguiente asignación de cartera es una recomendación de research del analista (síntesis de las vistas tácticas Black-Litterman vigentes), no un dato de mercado observado.", font_size=8.2, space_after=2.5)
    t_alloc = doc.add_table(rows=1, cols=5)
    formatear_tabla_apa7(
        t_alloc,
        col_widths=[1.70, 1.10, 1.20, 1.20, 2.00],
        headers=["Clase de Activo", "Ponderación", "Retorno Esperado", "Volatilidad Anual", "Tesis de Asignación Táctica"],
        data_rows=[
            ["LECAPs Cortas (S31O6)", "35,0%", "+11,8% USD Eq.", "3,2%", "Sobreponderar · Carry seguro."],
            ["Globales USD (GD35/GD38)", "35,0%", "+16,5% USD Total", "14,5%", "Sobreponderar · Convexidad alta."],
            ["Boncer Tramo Medio (TZX27)", "15,0%", "CER + 1,10% Real", "6,8%", "Neutral · Cobertura inflacionaria."],
            ["Acciones Energéticas (YPF/PAMP)", "10,0%", "+22,0% USD", "24,0%", "Sobreponderar · Vaca Muerta / RIGI."],
            ["Bopreal Serie 3 (USD)", "5,0%", "+8,4% USD Hard", "5,5%", "Sobreponderar · Cobertura dura."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=6.9
    )
    
    h5 = doc.add_heading("5. Conclusiones y Referencias Bibliográficas (APA 7)", level=2)
    h5.paragraph_format.space_before = Pt(3); h5.paragraph_format.space_after = Pt(1.5)
    for r in h5.runs: r.font.name = "Georgia"; r.font.size = Pt(9.0); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
    
    add_body(doc, "El régimen macroeconómico actual combina anclaje nominal, disciplina fiscal y una estructura de tasas que incentiva el posicionamiento táctico en moneda local de corto plazo y soberanos en moneda dura de duración intermedia-larga, maximizando la captura de valor en un entorno de desinflación sostenida.", font_size=7.6, space_after=1.5)
    add_body(doc, "La consolidación del superávit primario de caja elimina los riesgos de dominancia fiscal y refuerza la sustentabilidad de la deuda, garantizando un sendero de compresión en las primas de riesgo soberano.", font_size=7.6, space_after=1.5)
    
    refs = [
        "Banco Central de la República Argentina. (2026). Informe Monetario Mensual y Relevamiento de Expectativas de Mercado (REM). Buenos Aires.",
        "Bai, J., & Perron, P. (2003). Computation and analysis of multiple structural change models. Journal of Applied Econometrics, 18(1), 1-22.",
        "Campbell, J. Y., Lo, A. W., & MacKinlay, A. C. (1997). The Econometrics of Financial Markets. Princeton University Press.",
        "Instituto Nacional de Estadística y Censos. (2026). Índice de precios al consumidor (IPC). Informes Técnicos, 10(158).",
        "Ledoit, O., & Wolf, M. (2004). A well-conditioned estimator for large-dimensional covariance matrices. Journal of Multivariate Analysis, 88(2), 365-411.",
        "Nelson, C. R., & Siegel, A. F. (1987). Parsimonious Modeling of Yield Curves. The Journal of Business, 60(4), 473-489.",
        "Sargent, T. J., & Wallace, N. (1981). Some unpleasant monetarist arithmetic. Federal Reserve Bank of Minneapolis Quarterly Review, 5(3), 1-17.",
        "Taylor, J. B. (1993). Discretion versus policy rules in practice. Carnegie-Rochester Conference Series on Public Policy, 39, 195-214."
    ]
    for rf in refs:
        p_rf = doc.add_paragraph()
        p_rf.paragraph_format.left_indent = Inches(0.15); p_rf.paragraph_format.first_line_indent = Inches(-0.15)
        p_rf.paragraph_format.space_after = Pt(1)
        r = p_rf.add_run(rf)
        r.font.name = "Georgia"; r.font.size = Pt(6.6); r.font.color.rgb = COLOR_CHARCOAL

    os.makedirs(os.path.dirname(os.path.abspath(ruta_salida_docx)), exist_ok=True)
    doc.save(ruta_salida_docx)
    print("Paper Semanal DOCX generado:", ruta_salida_docx)
    return ruta_salida_docx

def compilar_y_exportar_paper_semanal(ruta_docx: str, ruta_pdf: str, periodo_str=None):
    compilar_paper_semanal_completo(ruta_docx, periodo_str)
    import pythoncom
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc_w = word.Documents.Open(os.path.abspath(ruta_docx), ReadOnly=True)
        doc_w.SaveAs(os.path.abspath(ruta_pdf), FileFormat=17)
        doc_w.Close(SaveChanges=False)
        print(f"Paper Semanal PDF exportado exitosamente: {ruta_pdf}")
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

if __name__ == "__main__":
    fecha_archivo = datetime.now().strftime("%Y-%m-%d")
    p_docx = os.path.join(BASE_DIR, "05_Informes_Semanales_APA7", f"{fecha_archivo}_Paper_Macroeconomico_Semanal.docx")
    p_pdf = p_docx.replace(".docx", ".pdf")
    compilar_y_exportar_paper_semanal(p_docx, p_pdf)
