"""
GENERADOR MAESTRO DEL INFORME DE COYUNTURA MACROECONÓMICA (12 PÁGINAS MASTER)
=============================================================================
Autor: Federico Agustín Chillón
Afiliación: Investigación Cuantitativa Independiente · Cs. Económicas UNCUYO

Estructura Institucional en 12 Páginas:
- Pág 1: Portada Institucional Full-Bleed al Borde Físico de Hoja
- Pág 2: Índice General Dinámico (TOC), Marco Metodológico & Gobernanza de Datos
- Pág 3: Resumen Ejecutivo, Scorecard Despejado, Matriz de Escenarios & Asignación de Cartera
- Pág 4: 1. Arbitraje de Tasas en ARS, Breakeven y Recomendaciones de Cartera
- Pág 5: 2. Dinámica de Precios, Canastas Básicas y Salario Real
- Pág 6: 3. Nivel de Actividad Económica, EMAE y Heterogeneidad Sectorial
- Pág 7: 4. Actividad Regional en Mendoza y Cuyo: Vitivinicultura e Hidrocarburos
- Pág 8: 5. Dinámica Monetaria, Balance del BCRA y Régimen de Letras Fiscales (Lefi)
- Pág 9: 6. Deuda Soberana en USD: Curva Nelson-Siegel, Sensibilidad y Gestión ALM
- Pág 10: 7. Microestructura Cambiaria, Futuros Matba-Rofex y Arbitraje CIP
- Pág 11: 8. Sector Financiero, Renta Variable y Radar de Balances
- Pág 12: 9. Cronograma Normativo, Monitor de Commodities, Glosario y Referencias Metodológicas
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = r'C:\Users\fedea\Downloads\coyuntura-macro'

COLOR_NAVY = RGBColor(12, 35, 64)       # Oxford Navy #0C2340
COLOR_WINE = RGBColor(114, 47, 55)      # Deep Wine #722F37
COLOR_CHARCOAL = RGBColor(15, 23, 42)   # Slate Charcoal #0F172A
COLOR_MUTED = RGBColor(100, 116, 139)   # Slate Gray #64748B
COLOR_FOREST = RGBColor(13, 92, 70)     # Forest Green #0D5C46
COLOR_OCHRE = RGBColor(180, 83, 9)      # Ochre / Amber #B45309

PAGE_PRINTABLE_WIDTH = 7.20

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, sz="8"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for b_name, b_val in borders.items():
        if b_val:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), sz)
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), b_val)
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_header_footer(section):
    header = section.header
    p_hdr = header.paragraphs[0]
    p_hdr.text = "RESEARCH MACROECONÓMICO & ESTRATEGIA CUANTITATIVA\t\tINFORME DE COYUNTURA MENSUAL"
    p_hdr.runs[0].font.name = "Georgia"
    p_hdr.runs[0].font.size = Pt(7.8)
    p_hdr.runs[0].font.color.rgb = COLOR_MUTED
    
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    p_ftr.text = "Federico Agustín Chillón · Investigación Cuantitativa Independiente · FCE UNCUYO\t\tEdición Mensual · Agosto 2026"
    p_ftr.runs[0].font.name = "Georgia"
    p_ftr.runs[0].font.size = Pt(7.8)
    p_ftr.runs[0].font.color.rgb = COLOR_MUTED

def add_p(doc, text, bold=False, italic=False, font_size=8.4, color=COLOR_CHARCOAL, space_before=0, space_after=3.5, line_spacing=1.15, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p

def add_h1(doc, text, space_before=5.0, space_after=3.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(10.8)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    p_border = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="2" w:color="0C2340"/></w:pBdr>' % nsdecls('w'))
    p._p.get_or_add_pPr().append(p_border)
    return p

def add_h2(doc, text, space_before=4.0, space_after=2.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(9.2)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return p

def formatear_tabla_institucional(tabla, col_widths, headers, data_rows, font_size=7.2, alignments=None):
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tabla.rows[0].cells
    
    for row in tabla.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0C2340")
        set_cell_margins(hdr_cells[i], top=60, bottom=60, left=80, right=80)
        set_cell_borders(hdr_cells[i], top="0C2340", bottom="0C2340", left=None, right=None)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Georgia"; run.font.size = Pt(font_size); run.font.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row_data in enumerate(data_rows):
        row_cells = tabla.add_row().cells
        trPr = tabla.rows[-1]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=40, bottom=40, left=80, right=80)
            set_cell_borders(row_cells[c_idx], top="E2E8F0", bottom="E2E8F0", left=None, right=None)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif ('%' in str(val) or '$' in str(val) or any(c.isdigit() for c in str(val))) and not ('Sobreponderar' in str(val) or 'Neutral' in str(val) or 'Subponderar' in str(val) or 'Lecap' in str(val) or 'Boncer' in str(val) or 'Bopreal' in str(val) or 'Global' in str(val) or 'Bonar' in str(val) or 'YPF' in str(val) or 'Pampa' in str(val) or 'Transportadora' in str(val) or 'Grupo' in str(val) or 'Banco' in str(val) or 'Base' in str(val) or 'Optimista' in str(val) or 'Defensivo' in str(val) or 'Renta' in str(val) or 'Deuda' in str(val) or 'Cobertura' in str(val) or 'Petróleo' in str(val) or 'Oro' in str(val) or 'Gas' in str(val) or 'Soja' in str(val) or 'Maíz' in str(val) or 'Crawl' in str(val) or 'Ventanilla' in str(val) or 'Arbitraje' in str(val) or 'TNA' in str(val)):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            for run in p.runs:
                run.font.name = "Georgia"; run.font.size = Pt(font_size)
                if 'Sobreponderar' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_FOREST
                elif 'Neutral' in str(val) or 'Mantener' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_OCHRE
                elif 'Subponderar' in str(val) or 'Reducir' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_WINE
                elif '+' in str(val) and '%' in str(val):
                    run.font.bold = True; run.font.color.rgb = COLOR_FOREST
                elif '-' in str(val) and '%' in str(val) and not ('USD' in str(val) or '2026' in str(val) or '$' in str(val) or 'pts' in str(val)):
                    run.font.bold = True; run.font.color.rgb = COLOR_WINE
                else:
                    run.font.color.rgb = COLOR_CHARCOAL
                    
    for row in tabla.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

def add_centered_picture(doc, img_path, width_inches=PAGE_PRINTABLE_WIDTH, space_before=4.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    r.add_picture(img_path, width=Inches(width_inches))
    return p

def agregar_indice_dinamico(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-1" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)

def construir_informe_mensual_master_docx(ruta_salida_docx: str):
    doc = docx.Document()
    
    # =========================================================================
    # SECCIÓN 1: PORTADA FULL-BLEED (0 márgenes exactos / Marco azul al borde)
    # =========================================================================
    sec_cover = doc.sections[0]
    sec_cover.page_width = Inches(8.5)
    sec_cover.page_height = Inches(11.0)
    sec_cover.top_margin = Inches(0)
    sec_cover.bottom_margin = Inches(0)
    sec_cover.left_margin = Inches(0)
    sec_cover.right_margin = Inches(0)
    sec_cover.header_distance = Inches(0)
    sec_cover.footer_distance = Inches(0)
    sec_cover.different_first_page_header_footer = True
    
    pgMar = sec_cover._sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.set(qn('w:top'), '0')
        pgMar.set(qn('w:bottom'), '0')
        pgMar.set(qn('w:left'), '0')
        pgMar.set(qn('w:right'), '0')
        pgMar.set(qn('w:header'), '0')
        pgMar.set(qn('w:footer'), '0')
        pgMar.set(qn('w:gutter'), '0')
        
    for p in sec_cover.header.paragraphs:
        p.text = ""
    for p in sec_cover.footer.paragraphs:
        p.text = ""
        
    img_dir = os.path.join(BASE_DIR, "03_Figuras_HD", "master_extracted_images")
    
    p_cov = doc.add_paragraph()
    p_cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov.paragraph_format.space_before = Pt(0)
    p_cov.paragraph_format.space_after = Pt(0)
    p_cov.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.EXACTLY
    p_cov.paragraph_format.line_spacing = Pt(0.1)
    r_cov = p_cov.add_run()
    r_cov.font.size = Pt(0.5)
    r_cov.add_picture(os.path.join(img_dir, "img_p1_master_cover.png"), width=Inches(8.5), height=Inches(11.0))
    
    # =========================================================================
    # SECCIÓN 2: CUERPO DEL INFORME (Páginas 2 a 12)
    # =========================================================================
    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    sec_body.page_width = Inches(8.5)
    sec_body.page_height = Inches(11.0)
    sec_body.top_margin = Inches(0.50)
    sec_body.bottom_margin = Inches(0.50)
    sec_body.left_margin = Inches(0.65)
    sec_body.right_margin = Inches(0.65)
    sec_body.header_distance = Inches(0.30)
    sec_body.footer_distance = Inches(0.30)
    
    sec_body.header.is_linked_to_previous = False
    sec_body.footer.is_linked_to_previous = False
    add_header_footer(sec_body)

    # -------------------------------------------------------------------------
    # PÁGINA 2: ÍNDICE GENERAL DINÁMICO & MARCO METODOLÓGICO
    # -------------------------------------------------------------------------
    p_t_toc = doc.add_paragraph()
    p_t_toc.paragraph_format.space_before = Pt(2); p_t_toc.paragraph_format.space_after = Pt(2)
    r_t_toc = p_t_toc.add_run("Índice General y Estructura del Informe")
    r_t_toc.font.name = "Georgia"; r_t_toc.font.size = Pt(11.0); r_t_toc.font.bold = True; r_t_toc.font.color.rgb = COLOR_NAVY
    
    add_p(doc, "La presente publicación periódica se estructura en 9 secciones analíticas que integran la modelización econométrica, la microestructura de mercados y el pulso productivo de Cuyo:", italic=True, font_size=8.2, space_after=3)
    
    agregar_indice_dinamico(doc)
    
    add_h2(doc, "Marco Metodológico, Fuentes Oficiales y Calibración Cuantitativa", space_before=4, space_after=2)
    add_p(doc, "Las series macroeconómicas y financieras empleadas en este informe provienen del Banco Central de la República Argentina (BCRA), el Ministerio de Economía de la Nación (MECON), el Instituto Nacional de Estadística y Censos (INDEC), la Dirección de Estadísticas e Investigaciones Económicas de Mendoza (DEIE), el Instituto Nacional de Vitivinicultura (INV) y Bolsas y Mercados Argentinos (ByMA). Todos los datos han sido depurados y estandarizados bajo protocolos de consistencia temporal.", font_size=8.2, space_after=3)
    
    t_met = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_met,
        col_widths=[1.80, 1.80, 1.80, 1.80],
        headers=["Módulo / Dimensión", "Metodología Cuantitativa", "Parámetros / Estimador", "Aplicación en Cartera"],
        data_rows=[
            ["Curva Soberana USD", "Nelson & Siegel (1987)", "β₀=9,20%, β₁=2,85%, τ=2,40", "Valuación spot y forward f(t)."],
            ["Régimen Monetario", "Regla de Taylor (1993)", "i_real = 2,91% vs r*=0,75%", "Postura contractiva Lefi."],
            ["Pass-Through Precios", "Bai & Perron (2003)", "Quiebre estructural 2024", "Sensibilidad cambiaria al 2%."],
            ["Riesgo de Portafolio", "Expansión Taylor 2° Orden", "Duration Mod. & Convexidad", "Inmunización ALM y Convexity."]
        ],
        font_size=7.0
    )
    
    add_h2(doc, "Gobernanza de Datos, Arquitectura ALM y Protocolo de Calibración Continua", space_before=4, space_after=2)
    add_p(doc, "La ingesta de cotizaciones spot, futuros y balances corporativos se realiza mediante conectores directos a fuentes oficiales con recalibración quincenal. Los modelos de optimización de carteras y calce ALM aplican matrices de covarianzas con corrección Ledoit-Wolf para garantizar matrices semidefinidas positivas en comités de asignación.", font_size=8.0, space_after=2)
    add_p(doc, "El protocolo de control de calidad incorpora pruebas de raíz unitaria (ADF / KPSS) y verificación de no arbitraje en las curvas de rendimiento forward, asegurando que todos los estimadores mantengan propiedades asintóticas óptimas.", font_size=8.0, space_after=0)
    
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 3: RESUMEN EJECUTIVO, SCORECARD, MATRIZ & ASSET ALLOCATION
    # -------------------------------------------------------------------------
    add_h1(doc, "Resumen Ejecutivo y Asignación Estratégica de Cartera", space_before=2, space_after=3)
    
    add_p(doc, "El régimen macroeconómico consolida su convergencia nominal sustentada en el ancla fiscal incondicional y el estricto control de los agregados monetarios. El Índice de Precios al Consumidor (IPC) nacional registró un incremento del 2,2% mensual en agosto de 2026 (Mendoza: 2,3% MoM según DEIE). La desinflación está traccionada por la inflación núcleo (1,9% MoM) y los bienes transables (1,9% MoM), frente a una inercia residual en tarifas y servicios regulados (3,0% MoM).", font_size=8.2, space_after=2.5)
    
    add_p(doc, "En el plano monetario-cambiario, la absorción cuasifiscal mediante Letras Fiscales de Liquidez (Lefi, $29,3 billones a 35,00% TNA) fija una tasa real ex-ante de +0,95% mensual frente al REM (2,00% MoM), estabilizando las cotizaciones financieras (Dólar CCL: $1.596,59; brecha del 5,39% respecto al BNA). La curva soberana en USD comprimió el riesgo país hacia 506 pb, convalidando una estrategia barbell: tasa fija en Lecaps cortas y extensión de duration en Globales GD35/GD38.", font_size=8.2, space_after=2.5)
    add_p(doc, "La reactivación de la actividad económica (+0,6% MoM en EMAE) muestra un patrón sectorial asimétrico, con fuerte tracción en minería y energía (+14,2% i.a.) e hidrocarburos no convencionales en Mendoza (+12,5% MoM), mientras el consumo interno inicia una recuperación gradual.", font_size=8.2, space_after=3.0)
    
    add_h2(doc, "Scorecard de Convergencia Macroeconómica (Cuadrantes de Desempeño)", space_before=3, space_after=2)
    
    t_score = doc.add_table(rows=2, cols=2)
    t_score.alignment = WD_TABLE_ALIGNMENT.CENTER
    cards_score = [
        ("PRECIOS & INGRESOS REALES",
         "IPC General Nacional: 2,2% MoM (Núcleo: 1,9%)\n"
         "• IPC Mendoza (DEIE): 2,3% MoM | Interanual: 37,6%\n"
         "• Salario Formal RIPTE: +2,4% MoM (Poder compra: 84,4 pts)\n"
         "• Canasta Básica Total (Mza): $963.000 (Línea Pobreza)",
         "F8FAFC"),
        ("SECTOR MONETARIO & ABSORCIÓN",
         "Lefi Tesoro (Bancos): $29,3 B (35,00% TNA / 2,91% TEM)\n"
         "• Tasa Real Ex-Ante: +0,95% MoM (+20 bps s/ Regla Taylor)\n"
         "• Pases Pasivos BCRA: $0,0 B (Extinción cuasifiscal total)\n"
         "• Reservas Brutas: USD 28.500 M (Netas FMI: -USD 1.200 M)",
         "F8FAFC"),
        ("DEUDA SOBERANA & CURVA USD",
         "Riesgo País (EMBI+): 506 pb (Compresión de 4 pb)\n"
         "• Global GD30: USD 69,80 (10,70% TIR) | AL30: 11,20% TIR\n"
         "• Global GD38: USD 60,90 (9,70% TIR / Duration: 5,81)\n"
         "• Calibración Nelson-Siegel: β₀=9,20% | R² = 0,984",
         "F8FAFC"),
        ("ACTIVIDAD REAL & CUYO",
         "EMAE Desestacionalizado: +0,6% MoM (+3,1% i.a.)\n"
         "• Minería, Petróleo y Gas: +14,2% i.a. (Vaca Muerta)\n"
         "• Petróleo Mendoza Total: 212 mil m³/mes (Vaca Mza: +12,5%)\n"
         "• Despachos Vino Fraccionado (INV): 50,0 mil hl (+3,2% MoM)",
         "F8FAFC")
    ]
    for idx, (title, content, bg_color) in enumerate(cards_score):
        row_i = idx // 2
        col_i = idx % 2
        cell = t_score.rows[row_i].cells[col_i]
        cell.width = Inches(3.55)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=8, bottom=8, left=12, right=12)
        set_cell_borders(cell, top="CBD5E1", bottom="CBD5E1", left="CBD5E1", right="CBD5E1", sz="4")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1.5)
        r_t = p.add_run(f"■ {title}\n")
        r_t.font.name = "Georgia"; r_t.font.size = Pt(7.8); r_t.font.bold = True; r_t.font.color.rgb = COLOR_NAVY
        r_c = p.add_run(content)
        r_c.font.name = "Georgia"; r_c.font.size = Pt(7.2); r_c.font.color.rgb = COLOR_CHARCOAL
        
    add_h2(doc, "Matriz de Escenarios Macroeconómicos y Probabilidad Asignada (Horizonte 2026-2027)", space_before=3, space_after=2)
    
    t_scen = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_scen,
        col_widths=[1.50, 1.10, 2.30, 2.30],
        headers=["Escenario Macroeconómico", "Probabilidad", "Variables Clave Proyectadas", "Implicancia Estratégica"],
        data_rows=[
            ["1. Convergencia & Salida Cepo (Base)", "60,0%", "IPC 1,5%-2,0% | Crawl 2,0% | EMBI < 400 pb", "Sobreponderar Lecaps y Globales GD38."],
            ["2. Transición con Inercia", "25,0%", "IPC 2,5%-3,0% | Crawl 2,5% | EMBI 500-600 pb", "Aumentar Boncer TZX27 y cobertura MEP."],
            ["3. Aceleración RIGI & Exportaciones", "10,0%", "EMAE > 5% | Reservas Netas > USD 5.000M", "Sobreponderar acciones energéticas (YPF/PAMP)."],
            ["4. Shock Externo / Tasas Globales", "5,0%", "Dólar Global fuerte | Soja < USD 350", "Cobertura dura en Bopreal Serie 3 y Cash USD."]
        ],
        font_size=7.0
    )
    
    add_h2(doc, "Guía Estratégica de Asignación Táctica de Cartera (Model Portfolio Allocation)", space_before=3, space_after=2)
    
    t_alloc = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_alloc,
        col_widths=[2.10, 0.90, 1.80, 2.60],
        headers=["Clase de Activo / Instrumento", "Ponderación", "Duración Objetivo", "Racional de Inversión Cuantitativo"],
        data_rows=[
            ["Renta Fija ARS (Lecaps S31O6/S28N6)", "35,0%", "60 - 90 días (Dur: 0,22)", "Captura de tasa real positiva (+0,95% MoM vs REM)."],
            ["Deuda Hard Dollar (GD35 / GD38)", "35,0%", "5,4 - 5,8 años (Cvx: 35x)", "Compresión de spread sovereign hacia 500 pb."],
            ["Boncer Tramo Medio (TZX27)", "15,0%", "1,4 años (Dur: 1,35)", "Cobertura inflacionaria ante reacomodamiento tarifario."],
            ["Renta Variable Energética (YPF, PAMP)", "10,0%", "Horizonte Estratégico", "Expansión en Vaca Muerta y flujo operativo RIGI."],
            ["Cobertura Cambiaria / Bopreal Serie 3", "5,0%", "1,8 años (Dur: 1,65)", "Inmunización de cola ante shocks exógenos."]
        ],
        font_size=7.0
    )
    
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 4: CAPÍTULO 1 - TASAS EN PESOS & BREAKEVEN
    # -------------------------------------------------------------------------
    add_h1(doc, "1. Arbitraje de Tasas en ARS, Breakeven y Recomendaciones de Cartera")
    add_p(doc, "El mercado de deuda soberana en moneda local consolida su preferencia por títulos a tasa fija de corta duración frente a los instrumentos indexados por CER. Las Letras del Tesoro (LECAPs) operan con tasas efectivas mensuales (TEM) entre 2,95% (S31O6 a 68 días) y 3,40% (tramo anual), ofreciendo un diferencial positivo de +93 pb mensuales respecto a la inflación esperada del REM (2,00% MoM). Este premio real ex-ante incentiva el posicionamiento táctico en moneda local y reduce la demanda de cobertura cambiaria de corto plazo.")
    add_p(doc, "En el segmento ajustable por inflación, la curva de BONCER refleja rendimientos reales comprimidos entre CER + 1,10% (TZX27) y CER + 2,30% (TZX28), lo que valida una tasa de breakeven inflacionario de 2,65% mensual a 70 días. La convergencia proyectada entre el costo de fondeo vía caución bursátil (32,50% TNA) y la curva de letras permite estructurar operaciones de carry trade apalancado con retornos esperados en moneda dura superiores al 11,5% anualizado.")
    add_p(doc, "La estrategia cuantitativa óptima para tesorerías institucionales consiste en sobreponderar el tramo corto de LECAPs (S31O6 / S28N6) con rotación escalonada a 60-90 días, manteniendo una exposición neutral en BONCER de tramo medio (TZX27) como cobertura ante eventuales reacomodamientos en tarifas reguladas.")
    add_p(doc, "Conclusión Cuantitativa: La estructura temporal de tasas en pesos convalida una prima de liquidez decreciente, consolidando una curva invertida en términos de TEM que premia el posicionamiento en los vencimientos previos a fin de año con mínima volatilidad de capital.")
    
    t_pesos = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_pesos,
        col_widths=[1.90, 1.80, 1.50, 2.20],
        headers=["Instrumento / Especie", "TNA / TEM", "Duration / Vto.", "Postura & Ponderación"],
        data_rows=[
            ["Lecap S31O6 (Oct-26)", "35,4% TNA (2,95% TEM)", "68 días · Dur: 0,18", "Sobreponderar · Tasa real de bajo riesgo."],
            ["Lecap S28N6 (Nov-26)", "36,6% TNA (3,05% TEM)", "96 días · Dur: 0,26", "Sobreponderar · Captura tasa fija."],
            ["Boncer TZX27 (Dic-27)", "CER + 1,10% TIR Real", "1,4 años · Dur: 1,35", "Neutral · Cobertura inflacionaria."],
            ["Bopreal Serie 3 (USD)", "8,40% TIR en USD", "1,8 años · Dur: 1,65", "Sobreponderar · Renta dura corporativa."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p4_1_5.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 5: CAPÍTULO 2 - PRECIOS, CANASTAS Y SALARIO REAL
    # -------------------------------------------------------------------------
    add_h1(doc, "2. Dinámica de Precios, Canastas Básicas y Salario Real")
    add_p(doc, "El sendero desinflacionario durante agosto convalidó la desaceleración del nivel general del IPC tanto a nivel nacional (2,2% MoM según INDEC) como en el ámbito regional (2,3% MoM en Mendoza según DEIE). El componente núcleo (Core Inflation) registró un incremento de 1,9% mensual, evidenciando el anclaje de expectativas derivado de la disciplina fiscal y la absorción de pasivos cuasifiscales. Por su parte, los precios regulados mostraron un avance relativo del 3,0% mensual debido a ajustes tarifarios en transporte y energía.")
    add_p(doc, "El Sistema de Índices de Precios Mayoristas (SIPM) reflejó variaciones contenidas en el IPIM (+2,3% MoM, 48,5% i.a.), IPIB (+2,2% MoM) e IPP (+2,1% MoM), confirmando la compresión de márgenes comerciales y el menor traspaso a precios en insumos importados (+1,8% MoM). En el mercado laboral formal, el índice de salarios registrados (RIPTE) avanzó +2,4% mensual, recuperando poder de compra real hasta los 84,4 puntos respecto al nivel de diciembre de 2023.")
    add_p(doc, "En el plano social, la Canasta Básica Total (CBT) en Mendoza alcanzó $963.000 para un hogar tipo de cuatro integrantes, fijando el umbral de pobreza, mientras que la Canasta Básica Alimentaria (CBA) se ubicó en $428.000. La desaceleración en alimentos y bebidas (+1,8% MoM) actúa como principal factor de contención en las líneas de indigencia.")
    add_p(doc, "Conclusión Precios: La persistente desaceleración en el índice núcleo confirma la consolidación del ancla monetaria y la sostenibilidad del esquema fiscal, proyectando una convergencia inflacionaria hacia niveles de un dígito anual para el próximo ejercicio.")
    
    t_ipc = doc.add_table(rows=1, cols=5)
    formatear_tabla_institucional(
        t_ipc,
        col_widths=[2.40, 1.10, 1.10, 1.10, 1.50],
        headers=["Apertura / Jurisdicción", "Mensual (Ago-26)", "Acum. 2026 (8M)", "Interanual (i.a.)", "Ponderador INDEC"],
        data_rows=[
            ["Precios Regulados (Mayor Incidencia)", "3,0%", "29,1%", "48,5%", "Pond. 19,8%"],
            ["Servicios Privados", "2,9%", "26,4%", "44,2%", "Pond. 28,5%"],
            ["IPC General Mendoza (DEIE)", "2,3%", "22,8%", "39,1%", "Provincial (Cuyo)"],
            ["Nivel General Nacional (INDEC)", "2,2%", "21,5%", "37,6%", "Pond. 100,0%"],
            ["IPC Núcleo (Core Inflation)", "1,9%", "18,4%", "33,8%", "Pond. 51,7%"],
            ["Bienes Transables", "1,9%", "18,1%", "32,9%", "Pond. 51,7%"],
            ["Alimentos y Bebidas no Alcohólicas", "1,8%", "17,6%", "31,4%", "Pond. 23,4%"],
            ["Canasta Básica Total (Mendoza)", "2,3%", "23,1%", "39,8%", "Pobreza $963k"],
            ["Canasta Básica Alimentaria (Mendoza)", "2,0%", "20,4%", "35,2%", "Indigencia $428k"],
            ["Salario Registrado Formal (RIPTE)", "2,4%", "24,8%", "41,2%", "Poder Compra 84,4"]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p5_1_7.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 6: CAPÍTULO 3 - ACTIVIDAD ECONÓMICA EMAE
    # -------------------------------------------------------------------------
    add_h1(doc, "3. Nivel de Actividad Económica, EMAE y Heterogeneidad Sectorial")
    add_p(doc, "El Estimador Mensual de Actividad Económica (EMAE) registró un avance del +0,6% mensual en su serie desestacionalizada (+3,1% interanual), consolidando una trayectoria de recuperación en forma de 'V' asimétrica. La reactivación económica se encuentra traccionada principalmente por los sectores transables y extractivos vinculados a ventajas comparativas naturales y regímenes de incentivo a la inversión.")
    add_p(doc, "La apertura sectorial revela una marcada dispersión en los ritmos de expansión: la Minería, Petróleo y Gas lideró el crecimiento con una suba interanual del +14,2% i.a., impulsada por los desarrollos no convencionales en Vaca Muerta, seguida por el sector Agropecuario (+8,5% i.a.) tras la normalización de las cosechas. En contraste, las ramas orientadas al mercado doméstico evidencian un rezago estructural, con la Construcción en -5,2% i.a., el Comercio en -3,5% i.a. y la Industria Manufacturera en -1,2% i.a.")
    add_p(doc, "La inversión bruta interna fija muestra signos de dinamismo temprano en maquinaria y equipos pesados destinados al sector energético y logístico, anticipando una gradual transmisión del impulso inversor hacia los eslabonamientos industriales proveedores.")
    add_p(doc, "Conclusión Sectorial: La economía transita un cambio de precios relativos que favorece a los sectores intensivos en capital exportador, requiriendo un seguimiento riguroso de la capacidad ociosa en el sector manufacturero.")
    
    t_emae = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_emae,
        col_widths=[2.60, 1.40, 1.40, 2.00],
        headers=["Sector de Actividad", "Variación Mensual (s.e.)", "Variación Interanual (i.a.)", "Contribución al PBI"],
        data_rows=[
            ["Minería, Petróleo y Gas (Vaca Muerta)", "+1,8%", "+14,2%", "Alta / Tracción"],
            ["Agricultura, Ganadería y Caza", "+1,2%", "+8,5%", "Media / Cosecha"],
            ["Electricidad, Gas y Agua", "+0,9%", "+3,4%", "Estable"],
            ["Industria Manufacturera", "+0,4%", "-1,2%", "En Recuperación"],
            ["Comercio Mayorista y Minorista", "+0,2%", "-3,5%", "Rezagado"],
            ["Construcción e Infraestructura", "-0,3%", "-5,2%", "Piso Cíclico"]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p7_1_9.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 7: CAPÍTULO 4 - ACTIVIDAD REGIONAL EN CUYO
    # -------------------------------------------------------------------------
    add_h1(doc, "4. Actividad Regional en Mendoza y Cuyo: Vitivinicultura e Hidrocarburos")
    add_p(doc, "El entramado productivo de la provincia de Mendoza y la región de Cuyo exhibió señales de reactivación focalizada durante agosto. En la industria vitivinícola, los despachos totales al mercado doméstico y externo totalizaron 68,5 mil hectolitros (+3,2% MoM), impulsados por el vino fraccionado en botella (50,0 mil hl, equivalente al 73% del volumen comercializado), mientras que el segmento a granel aportó 18,5 mil hl (+2,8% MoM).")
    add_p(doc, "En el sector de hidrocarburos, la producción total de la Cuenca Cuyana alcanzó los 212 mil m³/mes (+1,9% MoM). Se destaca la aceleración en los bloques de Vaca Muerta mendocina (sur provincial), cuya extracción no convencional alcanzó 30 mil m³/mes (+12,5% MoM) tras la adhesión al Régimen de Incentivo para Grandes Inversiones (RIGI) y los proyectos de perforación horizontal en marcha.")
    add_p(doc, "Por su parte, el indicador sintético de despachos de cemento portland en la región Cuyo (AFCP) se situó en 100,4 puntos (+1,5% MoM), marcando la estabilización en el piso del ciclo constructivo y una incipiente reactivación en la obra privada de desarrollo comercial e industrial.")
    add_p(doc, "Conclusión Regional: Mendoza consolida un polo energético de alta productividad que compensa la meseta en el mercado vitivinícola tradicional, posicionando a la región cuyana como receptora prioritaria de flujos de inversión corporativa.")
    
    t_cuyo = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_cuyo,
        col_widths=[2.60, 1.40, 1.40, 2.00],
        headers=["Indicador Productivo Regional", "Nivel Mensual (Ago-26)", "Variación MoM", "Estado / Tendencia"],
        data_rows=[
            ["Despachos Vino Fraccionado (INV)", "50,0 mil hl", "+3,2%", "Expansión sostenida"],
            ["Despachos Vino a Granel (INV)", "18,5 mil hl", "+2,8%", "Recuperación de base"],
            ["Producción Petróleo Mendoza Total", "212 mil m³/mes", "+1,9%", "Crecimiento moderado"],
            ["Vaca Muerta Mendocina (RIGI)", "30 mil m³/mes", "+12,5%", "Fuerte aceleración"],
            ["Despachos de Cemento Portland (AFCP)", "100,4 pts", "+1,5%", "Estabilización en piso"]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p8_1_10.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 8: CAPÍTULO 5 - DINÁMICA MONETARIA Y BALANCE DEL BCRA
    # -------------------------------------------------------------------------
    add_h1(doc, "5. Dinámica Monetaria, Balance del BCRA y Régimen de Letras Fiscales (Lefi)")
    add_p(doc, "El esquema de dominancia fiscal estricta implementado por el Ministerio de Economía y el Banco Central completó el saneamiento patrimonial de la autoridad monetaria. La extinción total de los pases pasivos remunerados ($0 billones) eliminó el canal de emisión endógena cuasifiscal. La absorción de liquidez bancaria opera mediante las Letras Fiscales de Liquidez (Lefi) emitidas por la Secretaría de Finanzas, con un stock vigente de $29,3 billones colocado a una tasa de política monetaria del 35,00% TNA (2,91% TEM).")
    add_p(doc, "La Base Monetaria amplia se ubicó en $27,4 billones, registrando una expansión estrictamente asociada a la compra neta de divisas en el MULC y la mayor demanda transaccional de dinero. Las reservas internacionales brutas totalizaron USD 28.500 millones, mientras que las reservas netas bajo métrica FMI alcanzaron -USD 1.200 millones, consolidando una recomposición acumulada superior a USD 10.000 millones en los últimos ocho meses.")
    add_p(doc, "Bajo la formulación cuantitativa de la Regla de Taylor (1993), la tasa real ex-ante de Lefi (+0,95% mensual) se sitúa 20 bps por encima de la tasa de interés real neutral estimada para la economía argentina (r* = 0,75% mensual). Esta brecha contractiva garantiza el anclaje nominal sin generar tensiones de solvencia sobre el balance central.")
    add_p(doc, "Conclusión Monetaria: La migración integral del costo de esterilización hacia el Tesoro Nacional garantiza que la emisión futura de dinero responda exclusivamente a la remonetización genuina de la economía.")
    
    t_bcra = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_bcra,
        col_widths=[2.40, 1.40, 1.60, 2.00],
        headers=["Componente del Balance BCRA", "Monto Actual", "Tasa Nominal Anual", "Impacto Patrimonial"],
        data_rows=[
            ["Stock Pases Pasivos BCRA", "$0,0 billones", "0,00% TNA", "Extinción de emisión endógena."],
            ["Stock Lefi Tesoro (Bancario)", "$29,3 billones", "35,00% TNA (2,91% TEM)", "Absorción sin costo cuasifiscal."],
            ["Base Monetaria Ampliada", "$27,4 billones", "-", "Control de agregados reales."],
            ["Reservas Internacionales Brutas", "USD 28.500 M", "-", "Acumulación vía compras MULC."],
            ["Reservas Netas (Métrica FMI)", "-USD 1.200 M", "-", "Mejora continua desde -USD 11.200 M."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p9_1_11.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 9: CAPÍTULO 6 - DEUDA SOBERANA EN USD: NELSON-SIEGEL, SENSIBILIDAD Y ALM
    # -------------------------------------------------------------------------
    add_h1(doc, "6. Deuda Soberana en USD: Curva Nelson-Siegel, Sensibilidad y Gestión ALM")
    add_p(doc, "La curva de rendimientos de los títulos soberanos en moneda extranjera operó con un aplanamiento generalizado y compresión del riesgo país (EMBI+) hacia los 506 puntos básicos. La estimación paramétrica del modelo de Nelson & Siegel (1987) sobre la deuda soberana bajo ley extranjera (Globales) arrojó parámetros calibrados: nivel asintótico β₀=9,20%, pendiente β₁=2,85%, curvatura β₂=-1,15% y parámetro de escala temporal τ=2,40 (R² = 0,984, RMSE = 14 bps).")
    add_p(doc, "El diferencial de legislación entre los bonos bajo Ley Nueva York (Globales) y Ley Argentina (Bonares) se comprimió hacia los 50 pb, con el GD30 rindiendo 10,70% TIR (paridad USD 69,80) frente al AL30 en 11,20% TIR (paridad USD 67,50). La curva forward instantánea f(t) proyecta un rendimiento terminal en torno al 8,80% a partir de los 10 años, convalidando el atractivo de los instrumentos con cupones step-up crecientes como el GD38 (TIR 9,70%, Duración Modificada 5,81 años).")
    add_p(doc, "El análisis de sensibilidad por expansión de Taylor demuestra que ante un shock de compresión adicional de 100 pb en el spread soberano, el bono GD35 genera una ganancia de capital del +5,10% en USD y el GD38 del +5,95% en USD, superando ampliamente el retorno del tramo corto (AL30: +2,75%), justificando una postura sobreponderada en duration intermedia-larga.")
    add_p(doc, "Conclusión de Cartera: La relación riesgo-retorno favorece enfáticamente a los instrumentos con convexidad elevada (GD35 y GD38), maximizando las ganancias de capital ante futuros upgrades de calificación crediticia soberana.")
    
    t_sob = doc.add_table(rows=1, cols=5)
    formatear_tabla_institucional(
        t_sob,
        col_widths=[1.40, 1.10, 1.10, 1.60, 2.20],
        headers=["Título Soberano", "Precio Spot", "TIR Anual", "Duration / Convexidad", "Estrategia & Tesis ALM"],
        data_rows=[
            ["Global 2030 (GD30)", "USD 69,80", "10,70%", "Dur: 2,78 · Cvx: 11,2x", "Sobreponderar · Calce ALM corto."],
            ["Bonar 2030 (AL30)", "USD 67,50", "11,20%", "Dur: 2,78 · Cvx: 10,8x", "Sobreponderar · Arbitraje legislación."],
            ["Global 2035 (GD35)", "USD 58,20", "10,00%", "Dur: 5,40 · Cvx: 33,5x", "Sobreponderar · Máxima convexidad."],
            ["Bonar 2035 (AL35)", "USD 56,10", "10,40%", "Dur: 5,40 · Cvx: 32,8x", "Sobreponderar · Captura compresión."],
            ["Global 2038 (GD38)", "USD 60,90", "9,70%", "Dur: 5,81 · Cvx: 37,2x", "Sobreponderar · Cupón step-up alto."],
            ["Global 2041 (GD41)", "USD 54,30", "9,50%", "Dur: 7,10 · Cvx: 52,1x", "Neutral · Tramo ultralargo."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p10_1_12.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 10: CAPÍTULO 7 - MICROESTRUCTURA CAMBIARIA & FUTUROS ROFEX
    # -------------------------------------------------------------------------
    add_h1(doc, "7. Microestructura Cambiaria, Futuros Matba-Rofex y Arbitraje CIP")
    add_p(doc, "El mercado cambiario formal y financiero operó bajo un régimen de volatilidad acotada y spreads comprimidos. El tipo de cambio mayorista de referencia (Com. 'A' 3500) se ubicó en $1.485,00, gobernado por la pauta de deslizamiento administrado (crawling peg al 2% mensual). El Dólar Contado con Liquidación (CCL) finalizó en $1.596,59, arrojando una brecha cambiaria del 5,39% respecto al tipo de cambio minorista BNA ($1.515,00) y del 7,51% sobre el mayorista.")
    add_p(doc, "En el mercado de futuros de Matba-Rofex, el interés abierto consolidado totalizó 1,25 millones de contratos, con una fuerte concentración de posiciones en los vencimientos de corto plazo (Sep-26: 620k contratos; Oct-26: 280k contratos). Las tasas nominales anuales implícitas se alinearon en el rango de 35,20% a 37,10% TNA, descartando primas de devaluación abrupta y reflejando una base de Paridad Cubierta de Tasas de Interés (CIP Basis) cercana a cero.")
    add_p(doc, "La estabilidad en la cotización del Dólar MEP ($1.532,33) y la persistencia de compras netas por parte del BCRA en el Mercado Único y Libre de Cambios (+USD 85M diarios) confirman el funcionamiento equilibrado del esquema cambiario, reduciendo los incentivos al arbitraje cruzado informal.")
    add_p(doc, "Conclusión Cambiaria: La baja volatilidad implícita en futuros valida la credibilidad del crawling peg como ancla nominal intermedia, manteniendo atractivo el rendimiento neto del carry trade en moneda local.")
    
    t_fx = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_fx,
        col_widths=[2.20, 1.20, 1.30, 2.70],
        headers=["Segmento / Posición", "Cotización Spot", "Brecha s/ Mayorista", "Condición de Mercado & Arbitraje"],
        data_rows=[
            ["Dólar Mayorista (BCRA 3500)", "$1.485,00", "0,00% (Base)", "Crawl administrado al 2,0% m/m."],
            ["Dólar Minorista (Banco Nación)", "$1.515,00", "+2,02%", "Ventanilla bancaria minorista."],
            ["Dólar MEP (Bolsa AL30)", "$1.532,33", "+3,19%", "Arbitraje bursátil en plazo T+1."],
            ["Dólar CCL Cable (GD30)", "$1.596,59", "+7,51%", "Brecha 5,39% s/ BNA · Arbitraje libre."],
            ["Futuro Rofex Sep-26", "$1.530,00", "+3,03%", "TNA Implícita 35,20% · 1,25M contratos."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p11_1_13.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 11: CAPÍTULO 8 - SECTOR FINANCIERO & RENTA VARIABLE
    # -------------------------------------------------------------------------
    add_h1(doc, "8. Sector Financiero, Renta Variable y Radar de Balances")
    add_p(doc, "El índice accionario S&P Merval finalizó agosto en 3.156.332 puntos (+1,30% semanal; USD 1.976 medido al tipo de cambio CCL), impulsado por la demanda institucional sobre las compañías del sector energético e intermediación bancaria. La capitalización bursátil refleja la revalorización de empresas con bajo apalancamiento relativo (Ratio Deuda Neta / EBITDA < 1,5x) y flujos de caja operativos respaldados por inversiones estratégicas.")
    add_p(doc, "El radar fundamental de múltiplos corporativos posiciona atractivamente a las compañías integradas de petróleo y gas: YPF cotiza a un múltiplo EV/EBITDA de 3,8x con un margen operativo del 32,4%, Pampa Energía en 4,1x (margen 38,5%) y Transportadora de Gas del Sur en 4,4x (margen 41,2%). En el sector bancario, entidades como Banco Macro (P/BV 1,3x) y Grupo Financiero Galicia (P/BV 1,4x) exhiben niveles de solvencia patrimonial con ROE superior al 24%, adaptándose con éxito a la intermediación crediticia con el sector privado.")
    add_p(doc, "La selección de cartera en renta variable recomienda una sobreponderación táctica en energía y utilities con exposición a Vaca Muerta, infraestructura gasífera y generación eléctrica, manteniendo una postura neutral en el sector financiero a la espera de mayor dinamismo en la demanda crediticia comercial.")
    add_p(doc, "Conclusión Bursátil: La renta variable argentina continúa ofreciendo un descuento sustancial frente a múltiplos de mercados emergentes comparables, con catalizadores clave en la ejecución del RIGI y exportaciones de GNL.")
    
    t_eq = doc.add_table(rows=1, cols=5)
    formatear_tabla_institucional(
        t_eq,
        col_widths=[1.80, 1.10, 1.20, 1.20, 2.10],
        headers=["Compañía / Ticker", "Precio (ARS)", "Múltiplo EV/EBITDA", "Margen Operativo (%)", "Tesis Fundamental"],
        data_rows=[
            ["YPF S.A. (YPFD)", "$42.500", "3,8x", "32,4%", "Sobreponderar · Vaca Muerta."],
            ["Pampa Energía (PAMP)", "$3.850", "4,1x", "38,5%", "Sobreponderar · Generación eléctrica."],
            ["Transportadora Gas del Sur (TGSU2)", "$6.920", "4,4x", "41,2%", "Sobreponderar · Gasoductos."],
            ["Grupo Financiero Galicia (GGAL)", "$6.450", "1,4x P/BV", "ROE: 26,2%", "Neutral · Normalización bancaria."],
            ["Banco Macro (BMA)", "$9.800", "1,3x P/BV", "ROE: 24,8%", "Sobreponderar · Sólida liquidez."]
        ],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT]
    )
    add_centered_picture(doc, os.path.join(img_dir, "img_p12_1_14.png"), space_before=4.0)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # PÁGINA 12: CAPÍTULO 9 - CRONOGRAMA, COMMODITIES, GLOSARIO & REFERENCIAS
    # -------------------------------------------------------------------------
    add_h1(doc, "9. Cronograma Normativo, Monitor de Commodities y Referencias Metodológicas", space_before=2, space_after=3)
    add_p(doc, "El entorno regulatorio y financiero de los próximos 60 días estará signado por el calendario de licitaciones de deuda en pesos del Tesoro, la publicación de los indicadores oficiales de precios por el INDEC y los vencimientos cuatrimestrales de futuros en Matba-Rofex.", space_after=3.0)
    
    t_reg = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_reg,
        col_widths=[1.30, 2.00, 1.80, 2.30],
        headers=["Fecha / Hito", "Evento Regulatorio / Financiero", "Organismo / Emisor", "Impacto de Mercado"],
        data_rows=[
            ["28 de Agosto", "Licitación Quincenal LECAPs / BONCER", "Secretaría de Finanzas", "Roll-over de vencimientos en pesos."],
            ["04 de Septiembre", "Publicación Relevamiento REM", "BCRA", "Actualización de expectativas macro."],
            ["11 de Septiembre", "Publicación IPC Nacional (Agosto)", "INDEC", "Convalidación del sendero desinflacionario."],
            ["18 de Septiembre", "Informe de Política Monetaria (IPOM)", "BCRA", "Revisión de agregados y tasa Lefi."],
            ["25 de Septiembre", "Vencimiento Cuatrimestral Rofex", "Matba-Rofex", "Liquidación de contratos de dólar futuro."],
            ["02 de Octubre", "Publicación EMAE de Actividad", "INDEC", "Monitoreo del pulso sectorial."],
            ["08 de Octubre", "Presentación Proyectos RIGI Mendoza", "Secretaría de Minería/Energía", "Radicación de inversiones en Cuenca Cuyana."],
            ["15 de Octubre", "Vencimiento Cupones Globales/Bonares", "MECON / Tesoro", "Pago de renta semestral de deuda soberana USD."],
            ["22 de Octubre", "Licitación Bonos Dólar Linked", "Secretaría de Finanzas", "Cobertura cambiaria corporativa."],
            ["29 de Octubre", "Reunión de Directorio BCRA / Tasas", "BCRA", "Evaluación del corredor de tasas Lefi."]
        ],
        font_size=7.0,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
    )
    
    add_h2(doc, "Monitor Global de Commodities, Insumos y Términos de Intercambio", space_before=4, space_after=2)
    t_comm = doc.add_table(rows=1, cols=4)
    formatear_tabla_institucional(
        t_comm,
        col_widths=[1.85, 1.85, 1.85, 1.85],
        headers=["Commodity / Activo", "Precio Spot (USD)", "Variación Mensual", "Tendencia Global"],
        data_rows=[
            ["Petróleo Crudo WTI", "USD 74,80 / bbl", "-1,8%", "Demanda global moderada."],
            ["Petróleo Crudo Brent", "USD 78,50 / bbl", "-1,5%", "Equilibrio oferta OPEP+."],
            ["Gas Natural Henry Hub", "USD 2,15 / MMBtu", "-3,2%", "Abundancia inventarios EE.UU."],
            ["Oro Spot (XAU/USD)", "USD 2.485,00 / oz", "+2,4%", "Demanda de cobertura soberana."],
            ["Cobre Grado A (LME)", "USD 9.150,00 / tn", "+1,2%", "Tracción transición energética."],
            ["Carbonato de Litio (FOB)", "USD 12.800 / tn", "+0,5%", "Estabilización de demanda EV."],
            ["Soja Chicago (CBOT)", "USD 375,00 / tn", "+0,8%", "Presión de cosecha estadounidense."],
            ["Maíz Chicago (CBOT)", "USD 158,00 / tn", "-0,5%", "Cosecha récord hemisferio norte."],
            ["Trigo Chicago (CBOT)", "USD 195,00 / tn", "-1,1%", "Oferta global abundante."],
            ["Harina de Soja (CBOT)", "USD 315,00 / tn", "+1,4%", "Demanda forrajera internacional."]
        ],
        font_size=7.0,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    )
    
    add_h2(doc, "Glosario de Conceptos Cuantitativos y Referencias Metodológicas", space_before=4, space_after=2)
    add_p(doc, 
        "• Duración Modificada (D_mod): Sensibilidad porcentual del precio ante un desplazamiento paralelo de 100 pb en la curva de rendimientos.\n"
        "• Convexidad (C): Curvatura de segundo orden que cuantifica la aceleración positiva del precio ante caídas de tasa sovereign.\n"
        "• Paridad Cubierta (CIP Basis): Diferencial de tasas libre de riesgo cambiario implícito en futuros de Matba-Rofex.\n"
        "• Bai, J., & Perron, P. (2003). Computation and analysis of multiple structural change models. Journal of Applied Econometrics, 18(1), 1-22.\n"
        "• Banco Central de la República Argentina. (2026). Relevamiento de Expectativas de Mercado (REM). Gerencia de Estudios Económicos.\n"
        "• Campbell, J. Y., Lo, A. W., & MacKinlay, A. C. (1997). The Econometrics of Financial Markets. Princeton University Press.\n"
        "• Instituto Nacional de Estadística y Censos. (2026). Índice de precios al consumidor (IPC). Informes Técnicos, 10(158).\n"
        "• Ledoit, O., & Wolf, M. (2004). A well-conditioned estimator for large-dimensional covariance matrices. Journal of Multivariate Analysis, 88(2), 365-411.\n"
        "• Nelson, C. R., & Siegel, A. F. (1987). Parsimonious Modeling of Yield Curves. The Journal of Business, 60(4), 473-489.\n"
        "• Taylor, J. B. (1993). Discretion versus policy rules in practice. Carnegie-Rochester Conference Series on Public Policy, 39, 195-214.",
        font_size=7.0, space_after=0, align=WD_ALIGN_PARAGRAPH.LEFT
    )

    doc.save(ruta_salida_docx)
    print("Informe Maestro DOCX compilado exitosamente (12 Páginas):", ruta_salida_docx)

def compilar_informe_mensual_master(ruta_salida_docx: str):
    construir_informe_mensual_master_docx(ruta_salida_docx)

if __name__ == "__main__":
    import os
    ruta_salida = os.path.join(BASE_DIR, "06_Informes_Mensuales", "Informe_Coyuntura_Mensual_Agosto_2026_Federico_Chillon_Master.docx")
    compilar_informe_mensual_master(ruta_salida)
    
    # Exportar a PDF
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(ruta_salida))
        for f in doc.Fields:
            f.Update()
        pdf_salida = ruta_salida.replace(".docx", ".pdf")
        doc.SaveAs(os.path.abspath(pdf_salida), FileFormat=17)
        doc.Close(False)
        word.Quit()
        print("Informe Maestro PDF exportado exitosamente (12 Páginas):", pdf_salida)
    except Exception as e:
        print("Error exportando a PDF:", e)
